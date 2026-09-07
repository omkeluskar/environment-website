#!/usr/bin/env python3
"""REFERENCES pages in the friend's college format, sourced from Chapters 1–7."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path("/workspace/cep-iot-bus-stop/chapters/References.docx")
TITLE = (
    "IoT-Based Bus Stop System for Public Transport Overcrowding "
    "Prediction Using Machine Learning"
)
TNR = "Times New Roman"


def set_run(run, *, size=12, bold=False, italic=False, color=None, underline=False):
    run.font.name = TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color is not None:
        run.font.color.rgb = color


def pfmt(p, *, align="left", before=0, after=0, line=16):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), TNR)
    rFonts.set(qn("w:hAnsi"), TNR)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading_label(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=10, after=2, line=16)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=12, bold=True)


def add_citation_parts(doc, parts):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=0, after=2, line=16)
    p.paragraph_format.keep_with_next = True
    for text, italic in parts:
        r = p.add_run(text)
        set_run(r, size=12, italic=italic)
    return p


def add_link_line(doc, prefix, url):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=0, after=6, line=16)
    r = p.add_run(f"{prefix} - ")
    set_run(r, size=12)
    add_hyperlink(p, url, url)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_page_border(section):
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), "000000")
        pgBorders.append(el)
    section._sectPr.append(pgBorders)


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    set_page_border(section)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    pfmt(hp, align="right", before=0, after=0, line=14)
    r = hp.add_run(TITLE)
    set_run(r, size=10, italic=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    pfmt(fp, align="center", before=0, after=0, line=14)
    add_page_field(fp)

    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "86")
    sectPr.append(pg)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = TNR
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    setup(doc)

    p = doc.add_paragraph()
    pfmt(p, align="center", before=6, after=14, line=20)
    r = p.add_run("REFERENCES")
    set_run(r, size=16, bold=True)

    add_heading_label(doc, "Research Paper 1")
    add_citation_parts(
        doc,
        [
            ("[1] T. Arabghalizi, A. Labrinidis, “Data-driven Bus Crowding Prediction Models Using Context-specific Features,” ", False),
            ("ACM/IMS Transactions on Data Science", True),
            (", vol. 1, no. 3, Art. no. 23, 2020.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.1145/3406962")

    add_heading_label(doc, "Research Paper 2")
    add_citation_parts(
        doc,
        [
            ("[2] M. Nitti, F. Pinna, L. Pintor, V. Pilloni, B. Barabino, “iABACUS: A Wi-Fi-Based Automatic Bus Passenger Counting System,” ", False),
            ("Energies", True),
            (", vol. 13, no. 6, Art. no. 1446, 2020.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.3390/en13061446")

    add_heading_label(doc, "Research Paper 3")
    add_citation_parts(
        doc,
        [
            ("[3] S. Tao, J. Corcoran, F. Rowe, M. Hickman, “To Travel or Not to Travel: ‘Weather’ Is the Question. Modelling the Effect of Local Weather Conditions on Bus Ridership,” ", False),
            ("Transportation Research Part C: Emerging Technologies", True),
            (", vol. 86, pp. 147–167, 2018.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.1016/j.trc.2017.11.005")

    add_heading_label(doc, "Research Paper 4")
    add_citation_parts(
        doc,
        [
            ("[4] J. Kniess, J. C. Rutke, W. A. C. Castaneda, “An IoT Transport Architecture for Passenger Counting: A Real Implementation,” ", False),
            ("2021 IFIP/IEEE International Symposium on Integrated Network Management (IM)", True),
            (", pp. 613–617, 2021.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://ieeexplore.ieee.org/document/9464024/")

    add_heading_label(doc, "Research Paper 5")
    add_citation_parts(
        doc,
        [
            ("[5] L. M. Pires, J. Figueiredo, R. Martins, J. Martins, “IoT-Enabled Real-Time Monitoring of Urban Garbage Levels Using Time-of-Flight Sensing Technology,” ", False),
            ("Sensors", True),
            (", vol. 25, no. 7, Art. no. 2152, 2025.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.3390/s25072152")

    doc.add_page_break()

    add_heading_label(doc, "Research Paper 6")
    add_citation_parts(
        doc,
        [
            ("[6] A. Olivo, G. Maternini, B. Barabino, “Empirical Study on the Accuracy and Precision of Automatic Passenger Counting in European Bus Services,” ", False),
            ("The Open Transportation Journal", True),
            (", vol. 13, pp. 250–260, 2019.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.2174/1874447801913010250")

    add_heading_label(doc, "Research Paper 7")
    add_citation_parts(
        doc,
        [
            ("[7] J. Wood, Z. Yu, V. V. Gayah, “Development and Evaluation of Frameworks for Real-Time Bus Passenger Occupancy Prediction,” ", False),
            ("International Journal of Transportation Science and Technology", True),
            (", vol. 12, no. 2, pp. 399–413, 2023.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.1016/j.ijtst.2022.03.005")

    add_heading_label(doc, "Research Paper 8")
    add_citation_parts(
        doc,
        [
            ("[8] E. Jenelius, “Personalized Predictive Public Transport Crowding Information with Automated Data Sources,” ", False),
            ("Transportation Research Part C: Emerging Technologies", True),
            (", vol. 117, Art. no. 102647, 2020.", False),
        ],
    )
    add_link_line(doc, "Official Website", "https://doi.org/10.1016/j.trc.2020.102647")

    add_heading_label(doc, "Book Reference –")
    add_citation_parts(
        doc,
        [
            ("[9] R. Kamal, ", False),
            ("Internet of Things: Architecture and Design Principles", True),
            (", 1st ed. New Delhi, India: McGraw Hill Education, 2017.", False),
        ],
    )
    add_link_line(
        doc,
        "Links",
        "https://books.google.co.in/books/about/Internet_of_Things.html?id=KuOizQEACAAJ&redir_esc=y",
    )

    add_heading_label(doc, "Book Reference –")
    add_citation_parts(
        doc,
        [
            ("[10] Transportation Research Board, ", False),
            ("Highway Capacity Manual", True),
            (", 6th ed. Washington, DC, USA: Transportation Research Board, 2016.", False),
        ],
    )
    add_link_line(doc, "Links", "https://www.trb.org/Main/Blurbs/175169.aspx")

    add_heading_label(doc, "YouTube Reference –")
    add_citation_parts(
        doc,
        [
            ("[11] Tech at Home, “Introduction to ESP32 - Getting Started | IoT Tutorials with ESP32 and Adafruit IO for Beginners,” YouTube. [Online Video].", False),
        ],
    )
    add_link_line(doc, "Links", "https://www.youtube.com/watch?v=dcMj8MWzJYk")

    add_heading_label(doc, "Community Survey Form –")
    add_citation_parts(
        doc,
        [
            ("[12] “Survey Form for IoT-Based Bus Stop System for Public Transport Overcrowding Prediction Using Machine Learning,” Google Forms. Conducted: Jul. 2026. [Online].", False),
        ],
    )
    p = doc.add_paragraph()
    pfmt(p, align="left", before=0, after=6, line=16)
    r = p.add_run("Available - ")
    set_run(r, size=12)
    r2 = p.add_run("(paste your 25-commuter survey Google Form link here)")
    set_run(r2, size=12, italic=True)

    doc.add_page_break()

    add_heading_label(doc, "Community Feedback Form –")
    add_citation_parts(
        doc,
        [
            ("[13] “Community Feedback Form for IoT-Based Bus Stop System for Public Transport Overcrowding Prediction Using Machine Learning,” Google Forms. Conducted: Jul. 2026. [Online].", False),
        ],
    )
    p = doc.add_paragraph()
    pfmt(p, align="left", before=0, after=6, line=16)
    r = p.add_run("Available - ")
    set_run(r, size=12)
    r2 = p.add_run("(paste your post-deployment feedback Google Form link here)")
    set_run(r2, size=12, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
