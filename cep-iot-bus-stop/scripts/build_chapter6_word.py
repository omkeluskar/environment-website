#!/usr/bin/env python3
"""Chapter 6 matching the student's PDF. Only extra mark: [13] in 6.6."""
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path("/workspace/cep-iot-bus-stop")
OUT_DOCX = ROOT / "chapters/Chapter_6_Community_Deployment_and_Impact_Assessment.docx"
FIG = ROOT / "figures/figure_6_1_satisfaction.png"
TITLE = (
    "IoT-Based Bus Stop System for Public Transport Overcrowding "
    "Prediction Using Machine Learning"
)
TNR = "Times New Roman"
NAVY = "1F4E79"


def set_run(run, *, size=12, bold=False, italic=False, font=TNR, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def pfmt(p, *, align="left", before=0, after=8, line=16):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


def add_title(doc, text, size=16, after=8, before=0):
    p = doc.add_paragraph()
    pfmt(p, align="center", before=before, after=after, line=18)
    r = p.add_run(text)
    set_run(r, size=size, bold=True)


def add_h(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=14, after=8, line=16)
    r = p.add_run(text)
    set_run(r, size=13, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="both", before=0, after=8, line=16)
    r = p.add_run(text)
    set_run(r, size=12)


def add_caption(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="center", before=8, after=6, line=14)
    r = p.add_run(text)
    set_run(r, size=11, bold=True)


def shade_cell(cell, fill=NAVY):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell(cell, text, *, header=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    pfmt(p, align="left", before=2, after=2, line=12)
    r = p.add_run(text)
    set_run(
        r,
        size=size,
        bold=header,
        color=RGBColor(255, 255, 255) if header else RGBColor(0, 0, 0),
    )


def add_table(doc, headers, rows, caption):
    add_caption(doc, caption)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        shade_cell(tbl.rows[0].cells[i])
        set_cell(tbl.rows[0].cells[i], h, header=True)
    for ri, row in enumerate(rows):
        fill = "F2F2F2" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            shade_cell(tbl.rows[ri + 1].cells[ci], fill)
            set_cell(tbl.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    pfmt(hp, align="center", before=0, after=0, line=12)
    r = hp.add_run(TITLE)
    set_run(r, size=10)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    pfmt(fp, align="center", before=0, after=0, line=12)
    add_page_field(fp)
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "72")
    sectPr.append(pg)


def make_figure():
    FIG.parent.mkdir(parents=True, exist_ok=True)
    labels = [
        "5 — Very\nSatisfied",
        "4 — Satisfied",
        "3 — Neutral",
        "2 — Dissatisfied",
        "1 — Very\nDissatisfied",
    ]
    values = [18, 5, 2, 0, 0]
    fig, ax = plt.subplots(figsize=(7.2, 3.6), dpi=160)
    bars = ax.bar(labels, values, color="#1F4E79", width=0.62)
    ax.set_ylabel("No. of Respondents")
    ax.set_ylim(0, 22)
    ax.set_title("Post-Deployment Satisfaction Rating Distribution")
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.3, str(v), ha="center", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(FIG, bbox_inches="tight", facecolor="white")
    plt.close()


def build():
    make_figure()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = TNR
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    setup_header_footer(doc)

    add_title(doc, "CHAPTER 6: COMMUNITY DEPLOYMENT AND", size=16, after=2, before=6)
    add_title(doc, "IMPACT ASSESSMENT", size=16, after=16)

    add_h(doc, "6.1 BENEFICIARY PROFILE")
    add_body(
        doc,
        "The system was developed for stakeholders connected with passenger movement at "
        "Dindoshi Bus Depot, BEST Route 326, Goregaon East, Mumbai. Its beneficiaries included "
        "commuters who required immediate on-site information, depot staff who needed remote "
        "operational updates, administrators responsible for transport decisions, field technicians "
        "maintaining the equipment, and academic evaluators assessing the feasibility of the prototype.",
    )
    add_body(
        doc,
        "Daily commuters were the primary on-site users. Their need for simple, glanceable crowding "
        "awareness without requiring a mobile application directly influenced the decision to use red "
        "and green LED indicators at the bus stop. Transit and depot staff served as direct remote "
        "users who needed clear risk indications to quickly judge severity and decide whether to "
        "dispatch feeder buses. Their need shaped the Telegram alert system and the Blynk dashboard "
        "interface.",
    )
    add_body(
        doc,
        "Depot administrators and decision-makers required a low-cost, low-maintenance solution with "
        "proactive alerts and no recurring subscription fees. Field technicians needed modular, "
        "easy-to-replace ESP32 hardware with simple I2C and GPIO connections on a 3.3 V rail. "
        "Academic evaluators required a fully documented system demonstrating correct application of "
        "IoT and machine learning concepts.",
    )
    add_body(
        doc,
        "The preference expressed by commuters for an on-site display directly influenced the use of "
        "red and green LEDs. This allowed passengers to interpret the current condition without "
        "installing an application or understanding the underlying risk score. Interviews with depot "
        "staff influenced the operator workflow. Staff rejected automatic passenger-count decay "
        "because it could reduce the count without evidence that a bus had departed. Consequently, "
        "the V7 button was used for controlled subtraction, while the V8 display showed the minutes "
        "since the last confirmed crossing before the operator applied the correction.",
    )

    add_h(doc, "6.2 COMMUNITY ENGAGEMENT ACTIVITIES")
    add_body(
        doc,
        "Community engagement was carried out before and during the July 2026 pilot deployment at "
        "the Dindoshi Bus Depot. The activities included direct observation, a structured commuter "
        "survey, informal interviews with depot staff, and coordination with the project guide to "
        "ensure the prototype met both academic and operational requirements.",
    )
    add_body(
        doc,
        "Routine observational visits were made to the Route 326 boarding edge during both morning "
        "and evening peak hours. During these visits, it was noted that commuters and transit staff "
        "had no objective method to assess crowding severity without being physically present. Crowd "
        "buildup occurred rapidly between 8:00 and 9:00 AM and between 6:00 and 7:00 PM, and manual "
        "intervention was entirely reactive, occurring only after overcrowding had become clearly "
        "visible. This observation confirmed the need for automated passenger sensing, continuous "
        "risk scoring, and early visual or remote alerts.",
    )
    add_body(
        doc,
        "A survey was conducted with 25 commuters who regularly use BEST services around the "
        "Dindoshi area. Among the respondents, 84% preferred an on-site physical display instead of "
        "a mobile application. This finding supported the decision to use physical LEDs at the bus "
        "stop rather than relying on an app-dependent design. A total of 92% wanted an overcrowding "
        "alert within 30 seconds of crowding buildup, indicating that delayed notifications would "
        "have limited operational value. In addition, 76% expected the passenger count to decrease "
        "gradually after a bus departed rather than being reset completely to zero. The most "
        "frequently reported frustration was the complete absence of any warning before crowding "
        "became excessive.",
    )
    add_body(
        doc,
        "Three depot staff members were interviewed informally during peak hours. They warned that "
        "automatic count decay could produce inaccurate values when passengers remained at the stop "
        "for extended periods, and they therefore preferred a human-controlled correction function. "
        "Staff also requested evidence before pressing any correction button, which led to the "
        "design of the V8 minutes-since-crossing display. Commuters similarly wanted visible signals "
        "that could be understood without technical knowledge, reinforcing the use of green for "
        "normal conditions and red for critical overcrowding.",
    )
    add_body(
        doc,
        "Coordination with the project guide ensured that the prototype met the academic CEP "
        "requirements while remaining practically relevant to the transit context. These combined "
        "engagement activities directly shaped the dual-interface design: simple LED indicators for "
        "commuters and a data-rich Blynk dashboard with Telegram alerts for operators.",
    )

    add_h(doc, "6.3 PILOT DEPLOYMENT")
    add_body(
        doc,
        "The pilot deployment was conducted at the Dindoshi Bus Depot, Route 326 boarding edge, "
        "during July 2026. Three demonstration sessions were completed using the integrated ESP32 "
        "hardware, Python backend, Blynk dashboard, Telegram alert channel, and CSV logging "
        "mechanism.",
    )
    add_table(
        doc,
        ["Location", "Date/Duration", "Participants", "Setup Details"],
        [
            [
                "Dindoshi Bus Depot, Route 326 boarding edge",
                "July 2026, 3 demonstration sessions",
                "1 project guide, 3 depot staff, 5 commuters",
                "ESP32 with VL53L0X + DHT22, laptop running brain.py, Blynk dashboard, Telegram bot, CSV logging cleared before each session",
            ],
        ],
        "Table 6.1: Pilot Deployment Details",
    )
    add_body(
        doc,
        "The risk score was calculated as the passenger count divided by the stop capacity of 45 "
        "and multiplied by 75. A rain bonus of 15, a heat bonus of 10 at temperatures of 33°C or "
        "above, and a rush-hour bonus of 10 during 8:00 to 9:00 AM and 6:00 to 7:00 PM were added "
        "when applicable. The score was capped at 98. Scores from 0 to 40 were classified as "
        "NORMAL, scores from 41 to 75 were classified as WARNING, and scores above 75 were "
        "classified as CRITICAL. Exactly 75 remained WARNING.",
    )
    add_body(
        doc,
        "In Scenario 1, the system was tested with 45 passengers, rain, a time of 17:30, and a "
        "temperature of 30°C The calculated risk was 90.0%, which was classified as CRITICAL. The "
        "red LED activated, and the Telegram alert was delivered in 0.81 seconds.",
    )
    add_body(
        doc,
        "In Scenario 2, the passenger count was reduced to 30 while rain, 17:30, and 30°C were "
        "maintained. The risk score became 65.0%, which was classified as WARNING. No Telegram "
        "alert was sent, and neither the critical red LED nor the normal green LED was activated.",
    )
    add_body(
        doc,
        "In Scenario 3, the passenger count was set to zero while rain remained active under "
        "conditions that did not qualify for heat or rush-hour bonuses. The resulting score was "
        "15.0%, which was classified as NORMAL, and the green LED activated.",
    )
    add_body(
        doc,
        "The operator function was also tested after a simulated bus departure. A staff member "
        "pressed V7 when the displayed passenger count was 50. The ESP32 applied max(0, passenger "
        "count minus 30), reducing the count from 50 to 20 without resetting it to zero. Every "
        "five-second brain.py cycle appended a row to bus_stop_log.csv throughout the pilot duration.",
    )
    add_body(
        doc,
        "The measured Blynk round-trip time was 891 ms, the Random Forest prediction required "
        "10.51 ms, and the CSV append operation was recorded as 0.00 ms at the available timing "
        "precision. The VL53L0X was polled every 150 ms, brain.py operated on a five-second cycle, "
        "and Telegram alerts used a 60-second cooldown.",
    )

    add_h(doc, "6.4 USER TRAINING")
    add_body(
        doc,
        "A five-minute walkthrough was conducted before participants interacted with the system. "
        "The demonstration introduced the ESP32 hardware, the Blynk dashboard, the Telegram "
        "notification channel, and the meaning of the physical LED indicators. Commuters were "
        "informed that the green LED represented a normal condition and that the red LED indicated "
        "critical overcrowding.",
    )
    add_body(
        doc,
        "Depot staff were trained to use the V7 Bus Left button only after confirming that a bus "
        "had departed. They were instructed to check the V8 minutes-since-crossing display first "
        "and then press V7 once for each bus departure. This procedure maintained human control "
        "and reduced the possibility of an unsupported count correction.",
    )
    add_body(
        doc,
        "One staff member initially pressed V7 twice in rapid succession. The mistake was handled "
        "immediately during training by repeating the one-press-per-departure instruction. The "
        "60-second cooldown applied to Telegram alerts and did not function as a lockout for the "
        "V7 button. One commuter initially confused the red LED with a fire alarm; however, the "
        "meaning of the indicator was clarified within seconds. These observations demonstrated "
        "the importance of brief operational instructions and clear on-device labels.",
    )

    add_h(doc, "6.5 FEEDBACK COLLECTION")
    add_body(
        doc,
        "Feedback was collected using a digital form and a physical notebook during in-person "
        "conversations. The two methods allowed participants to provide structured ratings as well "
        "as short comments about usability, trust, operational value, and desired improvements.",
    )
    add_body(
        doc,
        "Participants were asked how easy the system was to understand, whether they would trust "
        "the red and green LED indicators, how useful the Telegram alert was for their workflow, "
        "how they would rate the overall system, and what they would improve. The questions were "
        "kept short and plain so that all participants could respond without technical knowledge.",
    )
    add_body(
        doc,
        "The commuter responses focused mainly on the simplicity of the LED indicators and the "
        "value of receiving crowding information without an application. Staff feedback focused on "
        "the usefulness of Telegram alerts, the reliability of the passenger count, and the need "
        "for operator control when a bus departed. The collected responses were then grouped into "
        "common themes and summarised through average ratings.",
    )

    add_h(doc, "6.6 FEEDBACK ANALYSIS")
    add_body(
        doc,
        "The collected feedback [13] was systematically analysed to identify repeating patterns and "
        "common opinions across the participant groups. The most frequent positive comments "
        "highlighted the simplicity of the LED indicators and the value of receiving crowding "
        "information without requiring a mobile application. Participants appreciated the immediate "
        "visual clarity of the green and red states, which required no technical knowledge to "
        "interpret.",
    )
    add_body(
        doc,
        "The most common concerns related to the need for rush-hour validation and sensor "
        "direction. Several commuters noted that while the LED indicators were intuitive, they "
        "wanted to observe the system during actual peak-hour operation before fully trusting the "
        "readings. Depot staff raised the suggestion of adding a second sensor to distinguish "
        "passengers entering from those leaving, which would reduce the possibility of ghost counting.",
    )
    add_table(
        doc,
        ["Rating (out of 5)", "No. of Respondents", "Percentage (%)"],
        [
            ["5 — Very Satisfied", "18", "72.0%"],
            ["4 — Satisfied", "5", "20.0%"],
            ["3 — Neutral", "2", "8.0%"],
            ["2 — Dissatisfied", "0", "0.0%"],
            ["1 — Very Dissatisfied", "0", "0.0%"],
            ["Total", "25", "100%"],
        ],
        "Table 6.2: Post-Deployment Satisfaction Rating Distribution",
    )
    add_body(
        doc,
        "The satisfaction distribution showed that 92% of respondents rated the system as either "
        "satisfied or very satisfied. No participant rated the system below 3, indicating that the "
        "prototype was broadly well-received. The suggestion to add a second directional sensor "
        "was retained as a future improvement.",
    )
    add_caption(doc, "Figure 6.1: Bar Chart Satisfaction Rating Distribution")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG), width=Inches(6.3))

    add_h(doc, "6.7 IMPACT ASSESSMENT")
    add_body(
        doc,
        "The pilot demonstrated measurable changes in crowding awareness, alert speed, record "
        "keeping, operator control, and deployment cost. The system converted subjective visual "
        "estimation into a continuously updated risk score and provided both local and remote "
        "outputs suited to different beneficiary groups.",
    )
    add_table(
        doc,
        ["Indicator", "Before Deployment", "After Deployment", "Change"],
        [
            [
                "Crowding awareness",
                "No objective method; manual visual guesswork",
                "Continuous 5-second automated risk scoring",
                "Subjective to Objective",
            ],
            [
                "Alert response time",
                "Reactive; delay of minutes to hours",
                "0.81 seconds Telegram delivery on CRITICAL breach",
                "Minutes to Sub-second",
            ],
            [
                "Historical data",
                "No records; complaint-based only",
                "Continuous CSV logging every 5 seconds",
                "None to Persistent record",
            ],
            [
                "Operator workflow",
                "No standardised correction method",
                "V7 push button with V8 evidence display",
                "None to Human-in-the-loop control",
            ],
            [
                "Cost per stop",
                "High (industrial APC) or free but ineffective (manual)",
                "Under ₹1,500 (ESP32 hardware) + free APIs",
                "Prohibitive to Affordable",
            ],
        ],
        "Table 6.3: Impact Assessment",
    )
    add_body(
        doc,
        "The prototype remained financially accessible because its hardware cost was under ₹1,500 "
        "and its cloud, weather, and messaging services operated through free-tier APIs. This made "
        "the design appropriate for an academic pilot and potentially suitable for wider trials "
        "without an immediate subscription cost.",
    )

    add_h(doc, "6.8 OUTCOME ACHIEVEMENT MATRIX")
    add_body(
        doc,
        "The observed pilot results were mapped against the expected project outcomes. The matrix "
        "showed that the system achieved its main objectives relating to commuter awareness, staff "
        "notification, risk scoring, operator control, affordability, and historical data preservation.",
    )
    add_table(
        doc,
        ["Expected Outcome", "Achieved?", "Evidence"],
        [
            [
                "Commuters receive instant on-site crowding awareness",
                "Yes",
                "LED indicators successfully interpreted by commuters without training",
            ],
            [
                "Transit staff receive timely remote alerts",
                "Yes",
                "Telegram delivered in 0.81 s during CRITICAL scenario",
            ],
            [
                "System accurately predicts overcrowding risk",
                "Yes",
                "90.0% risk at full capacity with rain, 65.0% at 30 pax, 15.0% at 0 pax — matches formula exactly",
            ],
            [
                "Operator can safely correct passenger count",
                "Yes",
                "V7 reduced count from 50 to 20; V8 provided evidence before correction",
            ],
            [
                "System is affordable for municipal deployment",
                "Yes",
                "Hardware under ₹1,500, all APIs on free tier",
            ],
            [
                "Historical data is preserved for future analysis",
                "Yes",
                "bus_stop_log.csv appended every 5-second cycle",
            ],
        ],
        "Table 6.4: Outcome Achievement Matrix",
    )
    add_body(
        doc,
        "The three pilot scores matched the defined formula under the stated conditions. They also "
        "demonstrated a clear progression from NORMAL at zero passengers to WARNING at 30 "
        "passengers and CRITICAL at the stop capacity of 45 passengers when rain was present.",
    )

    add_h(doc, "6.9 SUSTAINABILITY PLAN")
    add_body(
        doc,
        "Operational sustainability was supported by the system's limited daily maintenance "
        "requirements. The 3.3 V rail design provided a consistent low-voltage supply to the "
        "connected modules, while the modular I2C and GPIO wiring allowed individual sensors or "
        "indicators to be replaced without rebuilding the complete circuit. Field technicians could "
        "therefore inspect and replace the VL53L0X, DHT22, or LED connections using standard "
        "electronic maintenance practices without specialised training.",
    )
    add_body(
        doc,
        "Technical sustainability was supported by the use of an open-source Python backend and "
        "Arduino firmware. The documented virtual-pin ownership, GPIO assignments, risk "
        "thresholds, and processing rules would allow future students or municipal "
        "information-technology staff to maintain and modify the prototype. Continuous accumulation "
        "of records in bus_stop_log.csv would also provide a foundation for future retraining of "
        "the machine-learning model using real-world observations instead of depending only on the "
        "initial formula-generated data.",
    )
    add_body(
        doc,
        "Future expansion could involve deploying additional ESP32 edge nodes at more bus stops "
        "while retaining the same layered architecture. A later version could migrate CSV storage "
        "to a cloud database such as MySQL or Firebase to support centralised multi-stop records. "
        "Solar-power integration could reduce dependence on fixed electrical connections, while an "
        "additional directional sensor could distinguish passengers entering from those leaving. "
        "These developments would improve the long-term viability of the system across Mumbai's "
        "public transport network.",
    )

    doc.save(OUT_DOCX)
    print("wrote", OUT_DOCX)


if __name__ == "__main__":
    build()
