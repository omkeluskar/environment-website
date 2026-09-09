#!/usr/bin/env python3
"""Logics appendix: VS Code-style screenshots of sanitized project source."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pygments import highlight
from pygments.formatters import ImageFormatter
from pygments.lexers import CppLexer, PythonLexer, TextLexer
from reportlab.lib.colors import black
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

ROOT = Path("/workspace/cep-iot-bus-stop")
SRC = ROOT / "source"
FIG = ROOT / "figures/logics"
CH = ROOT / "chapters"
TITLE = (
    "IoT-Based Bus Stop System for Public Transport Overcrowding "
    "Prediction Using Machine Learning"
)
TNRB = "Times-Bold"
TNR = "Times-Roman"
MONO = "DejaVu Sans Mono"

BRAIN_CORE = r'''import os, csv, time, requests, numpy as np
from datetime import datetime
from sklearn.ensemble import RandomForestRegressor

BLYNK_AUTH_TOKEN = "********"
BLYNK_URL = "https://blynk.cloud/external/api"
TELEGRAM_BOT_TOKEN = "********"
TELEGRAM_CHAT_ID = "********"
OPENWEATHER_API_KEY = "********"

BUS_STOP = {
    "id": "BEST-DIN-001",
    "name": "Dindoshi Bus Depot",
    "lat": 19.1753,
    "lon": 72.8649,
    "assigned_route": "Route 326",
}

STOP_CAPACITY = 45  # reasoned stand-capacity estimate

def calculate_synthetic_risk(hour, is_raining, passenger_count, temperature):
    passenger_ratio = min(1.0, passenger_count / STOP_CAPACITY)
    passenger_component = passenger_ratio * 75
    rain_bonus = 15 if is_raining else 0
    heat_bonus = 10 if temperature >= 33 else 0
    rush_hour_bonus = 10 if hour in [8, 9, 18, 19] else 0
    return min(98, passenger_component + rain_bonus + heat_bonus + rush_hour_bonus)

hours_to_sample = list(range(0, 24))
passenger_counts = list(range(0, 101))
weather_options = [0, 1]
temperature_options = [20, 24, 28, 30, 32, 33, 34, 35, 38]

X_train_list, y_train_list = [], []
for hour in hours_to_sample:
    for count in passenger_counts:
        for rain in weather_options:
            for temperature in temperature_options:
                risk = calculate_synthetic_risk(hour, rain, count, temperature)
                X_train_list.append([hour, rain, count, temperature])
                y_train_list.append(risk)

X_train, y_train = np.array(X_train_list), np.array(y_train_list)
model = RandomForestRegressor(n_estimators=100, random_state=42)
model.fit(X_train, y_train)
'''

BRAIN_LOOP = r'''def get_weather():
    try:
        url = (f"https://api.openweathermap.org/data/2.5/weather"
               f"?lat={BUS_STOP['lat']}&lon={BUS_STOP['lon']}"
               f"&appid={OPENWEATHER_API_KEY}&units=metric")
        data = requests.get(url, timeout=10).json()
        main = data["weather"][0]["main"]
        if main in ("Rain", "Drizzle", "Thunderstorm", "Shower"):
            return "Rain"
        return "Clear"
    except Exception:
        return "Clear"   # fallback: Clear, 28.0 C used if DHT22 missing

def get_blynk_value(pin):
    res = requests.get(
        f"{BLYNK_URL}/get?token={BLYNK_AUTH_TOKEN}&pin={pin}", timeout=5)
    val = res.text.strip('[]" \n')
    if "error" in val.lower() or not val:
        return None
    return float(val)

if not os.path.exists("bus_stop_log.csv"):
    with open("bus_stop_log.csv", "w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerow(
            ["timestamp", "passengers", "temp", "weather", "risk", "status"])

while True:
    weather_status = get_weather()
    is_raining = 1 if weather_status == "Rain" else 0
    passenger_count = int(float(get_blynk_value("V1") or 0))
    temp_value = float(get_blynk_value("V4") or 28.0)
    features = np.array([[datetime.now().hour, is_raining,
                          passenger_count, temp_value]])
    risk_score = model.predict(features)[0]

    # <=40 NORMAL | 41-75 WARNING | >75 CRITICAL (exactly 75 = WARNING)
    if risk_score > 75:
        csv_status = "CRITICAL"
        if time.time() - last_telegram_alert_time > 60:
            send_telegram_alert(BUS_STOP, passenger_count, risk_score, weather_status)
            last_telegram_alert_time = time.time()
    elif risk_score > 40:
        csv_status = "WARNING"
    else:
        csv_status = "NORMAL"

    with open("bus_stop_log.csv", "a", newline="", encoding="utf-8") as f:
        csv.writer(f).writerow([
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            passenger_count, f"{temp_value:.1f}", weather_status,
            f"{risk_score:.1f}", csv_status])
    time.sleep(5)
'''

BRAIN_TELEGRAM = r'''def send_telegram_alert(stop, passengers, risk, weather):
    message = (
        f"HIGH OVERCROWDING ALERT\n"
        f"Stop Name: {stop['name']}\n"
        f"Stop ID: {stop['id']}\n"
        f"Live Passenger Count: {passengers}\n"
        f"AI Risk Score: {risk:.1f}%\n"
        f"Weather: {weather}\n"
        f"Action: Dispatch extra feeder bus on {stop['assigned_route']}."
    )
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    payload = {"chat_id": TELEGRAM_CHAT_ID, "text": message}
    requests.post(url, json=payload, timeout=10)

# Blynk writes after each 5-second cycle (Python owns V2, V3, V5, V6)
requests.get(f"{BLYNK_URL}/update?token={BLYNK_AUTH_TOKEN}"
             f"&pin=V2&value={'Rain' if is_raining else 'Clear'}", timeout=5)
requests.get(f"{BLYNK_URL}/update?token={BLYNK_AUTH_TOKEN}"
             f"&pin=V3&value={int(risk_score)}", timeout=5)
requests.get(f"{BLYNK_URL}/update?token={BLYNK_AUTH_TOKEN}"
             f"&pin=V5&value={current_datetime_str}", timeout=5)
requests.get(f"{BLYNK_URL}/update?token={BLYNK_AUTH_TOKEN}"
             f"&pin=V6&value={status_emoji}", timeout=5)
'''

ARDUINO = r'''#define BLYNK_AUTH_TOKEN "********"
char ssid[] = "********";
char pass[] = "********";

#define DHTPIN 4
#define RED_LED 25
#define GREEN_LED 26
const int CONFIRM_THRESHOLD = 2;
const int DETECTION_THRESHOLD_MM = 400;
const unsigned long TOF_CHECK_INTERVAL = 150;

bool isPersonDetected() {
  VL53L0X_RangingMeasurementData_t measure;
  lox.rangingTest(&measure, false);
  return (measure.RangeStatus != 4 &&
          measure.RangeMilliMeter < DETECTION_THRESHOLD_MM);
}

void checkPassengerCrossing() {
  bool detected = isPersonDetected();
  if (detected) {
    consecutiveDetections++;
    if (consecutiveDetections >= CONFIRM_THRESHOLD && !personInFrame) {
      passengerCount++;
      personInFrame = true;
      lastCrossingTime = millis();
    }
  } else {
    consecutiveDetections = 0;
    personInFrame = false;
  }
}

// V7 Bus Left: subtract 30, never auto-reset to zero
BLYNK_WRITE(V7) {
  if (param.asInt() == 1) {
    passengerCount = max(0, passengerCount - 30);
  }
}

// V3 from Python: red LED only when risk > 75 (CRITICAL)
BLYNK_WRITE(V3) {
  pythonRiskScore = param.asInt();
  if (pythonRiskScore > 75) {
    digitalWrite(RED_LED, HIGH);
    digitalWrite(GREEN_LED, LOW);
  } else {
    digitalWrite(RED_LED, LOW);
    digitalWrite(GREEN_LED, HIGH);
  }
}

void sendDashboardData() {
  Blynk.virtualWrite(V1, passengerCount);
  Blynk.virtualWrite(V4, dht.readTemperature());
  Blynk.virtualWrite(V8, (millis() - lastCrossingTime) / 60000);
}
'''

CSV = '''timestamp,passengers,temp,weather,risk,status
2026-08-11 17:22:43,5,30.4,Clear,16.5,NORMAL
2026-08-11 17:51:42,30,30.3,Rain,65.0,WARNING
2026-08-11 17:56:44,0,30.3,Rain,15.0,NORMAL
2026-08-24 00:49:29,53,31.2,Rain,90.0,CRITICAL
2026-09-03 19:47:46,11,30.9,Rain,43.3,WARNING
2026-09-03 19:54:19,53,31.0,Rain,98.0,CRITICAL
2026-09-04 07:45:42,54,27.8,Rain,90.0,CRITICAL
2026-09-04 07:48:37,0,27.8,Rain,15.0,NORMAL
'''

PAGES = [
    ("1. BACKEND CORE LOGIC",
     "This code builds a 43,632-sample synthetic grid and trains RandomForestRegressor(n_estimators=100, random_state=42) on hour, rain, passenger count and temperature using the locked risk formula.",
     "brain.py", BRAIN_CORE, "python",
     "Image.1: Backend Core Logic"),
    ("2. LIVE MONITORING AND CSV LOG",
     "This code fetches Dindoshi weather from OpenWeatherMap, reads V1 and V4 from Blynk, predicts risk, classifies NORMAL / WARNING / CRITICAL, and appends one row to bus_stop_log.csv every five seconds.",
     "brain.py", BRAIN_LOOP, "python",
     "Image.2: Live Loop and CSV Logging"),
    ("3. TELEGRAM ALERT AND BLYNK WRITE",
     "This code sends a Telegram overcrowding alert when risk is above 75, subject to a 60-second cooldown, and writes weather, risk, time and status back to Blynk pins V2, V3, V5 and V6.",
     "brain.py", BRAIN_TELEGRAM, "python",
     "Image.3: Telegram and Blynk Output"),
    ("4. ESP32 FIRMWARE LOGIC",
     "This firmware confirms a passenger crossing after two consecutive VL53L0X readings below 400 mm, applies V7 as max(0, count-30), drives the LEDs from the Python risk on V3, and publishes V1, V4 and V8.",
     "smart_bus_stop.ino", ARDUINO, "cpp",
     "Image.4: ESP32 Crossing, V7 and LED Logic"),
    ("5. CSV SENSOR LOG",
     "This extract of bus_stop_log.csv shows live cycles at Dindoshi, including 65.0% WARNING at 30 passengers with rain and 90.0% CRITICAL at full-capacity rain cases.",
     "bus_stop_log.csv", CSV, "text",
     "Image.5: CSV Log Extract"),
]


def code_image(code, lexer_name, filename, out_path):
    lexer = {"python": PythonLexer(), "cpp": CppLexer(), "text": TextLexer()}[lexer_name]
    formatter = ImageFormatter(
        font_name=MONO,
        font_size=13,
        line_numbers=True,
        line_number_chars=3,
        line_number_bg="#1e1e1e",
        line_number_fg="#858585",
        line_number_separator=False,
        style="monokai",
        image_pad=10,
        line_pad=2,
    )
    raw = highlight(code.strip("\n") + "\n", lexer, formatter)
    tmp = out_path.with_suffix(".raw.png")
    tmp.write_bytes(raw)
    body = Image.open(tmp).convert("RGB")
    bar_h = 36
    pad = 8
    w = max(body.width + 2 * pad, 900)
    h = body.height + bar_h + 2 * pad
    img = Image.new("RGB", (w, h), "#2d2d2d")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w, bar_h], fill="#3c3c3c")
    for i, col in enumerate(["#ff5f56", "#ffbd2e", "#27c93f"]):
        draw.ellipse([10 + i * 16, 12, 22 + i * 16, 24], fill=col)
    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    draw.text((70, 10), filename, fill="#dddddd", font=font)
    img.paste(body, (pad, bar_h + pad))
    img.save(out_path)
    tmp.unlink(missing_ok=True)
    return out_path


def draw_pdf_page(c, heading, blurb, img_path, caption, page_no):
    W, H = letter
    margin = 43
    # header
    c.setFont(TNR, 10)
    c.drawCentredString(W / 2, H - 36, TITLE)
    c.setStrokeColor(black)
    c.setLineWidth(0.6)
    c.line(margin, H - 42, W - margin, H - 42)
    c.setFont(TNRB, 13)
    c.drawString(margin, H - 68, heading)
    c.setFont(TNR, 11)
    # wrap blurb
    words = blurb.split()
    lines, cur = [], ""
    maxw = W - 2 * margin
    for w in words:
        trial = (cur + " " + w).strip()
        if stringWidth(trial, TNR, 11) <= maxw:
            cur = trial
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    y = H - 88
    for line in lines:
        c.drawString(margin, y, line)
        y -= 14
    y -= 8
    im = Image.open(img_path)
    max_w = W - 2 * margin
    max_h = y - 70
    scale = min(max_w / im.width, max_h / im.height, 1.0)
    dw, dh = im.width * scale, im.height * scale
    c.drawImage(str(img_path), (W - dw) / 2, y - dh, width=dw, height=dh, preserveAspectRatio=True, mask="auto")
    c.setFont(TNRB, 11)
    c.drawCentredString(W / 2, y - dh - 18, caption)
    c.setFont(TNR, 11)
    c.drawCentredString(W / 2, 32, str(page_no))


def add_page_field(paragraph):
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_run(run, *, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def build_word(image_paths):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.0)
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(TITLE)
    set_run(r, size=10)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "89")
    sec._sectPr.append(pg)

    for i, (heading, blurb, _fn, _code, _lex, caption) in enumerate(PAGES):
        if i:
            doc.add_page_break()
        p = doc.add_paragraph()
        r = p.add_run(heading)
        set_run(r, size=14, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(blurb)
        set_run(r, size=12)
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(image_paths[i]), width=Inches(6.6))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run(r, size=11, bold=True)
    out = CH / "Appendix_A_System_Logic.docx"
    doc.save(out)
    print("wrote", out)


def main():
    FIG.mkdir(parents=True, exist_ok=True)
    SRC.mkdir(parents=True, exist_ok=True)
    (SRC / "brain_core.py").write_text(BRAIN_CORE)
    (SRC / "brain_loop.py").write_text(BRAIN_LOOP)
    (SRC / "brain_telegram.py").write_text(BRAIN_TELEGRAM)
    (SRC / "smart_bus_stop.ino").write_text(ARDUINO)
    (SRC / "bus_stop_log_extract.csv").write_text(CSV)
    (SRC / "README.md").write_text(
        "Sanitized excerpts for the black-book logics appendix.\n"
        "Live Blynk, Telegram, OpenWeatherMap and Wi-Fi secrets were replaced with ********.\n"
    )

    images = []
    for i, (_h, _b, filename, code, lexer, _c) in enumerate(PAGES, 1):
        path = FIG / f"image_{i}_{filename.replace('.', '_')}.png"
        code_image(code, lexer, filename, path)
        images.append(path)
        print("shot", path)

    pdf_path = CH / "Appendix_A_System_Logic.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    for i, page in enumerate(PAGES):
        draw_pdf_page(c, page[0], page[1], images[i], page[5], 89 + i)
        c.showPage()
    c.save()
    print("wrote", pdf_path)
    build_word(images)


if __name__ == "__main__":
    main()
