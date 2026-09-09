#!/usr/bin/env python3
"""Chapter 7 Word from the uploaded PDF. Header + PAGE start 80."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(
    "/workspace/cep-iot-bus-stop/chapters/"
    "Chapter_7_Results_Discussion_Conclusion_and_Future_Scope.docx"
)
TITLE = (
    "IoT-Based Bus Stop System for Public Transport Overcrowding "
    "Prediction Using Machine Learning"
)
TNR = "Times New Roman"


def set_run(run, *, size=12, bold=False, italic=False, font=TNR, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def pfmt(p, *, align="left", before=0, after=8, line=16):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


def add_title(doc, text, size=16, after=8, before=0):
    p = doc.add_paragraph()
    pfmt(p, align="center", before=before, after=after, line=18)
    r = p.add_run(text)
    set_run(r, size=size, bold=True)


def add_h(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=14, after=8, line=16)
    r = p.add_run(text)
    set_run(r, size=13, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="both", before=0, after=8, line=16)
    r = p.add_run(text)
    set_run(r, size=12)


def add_bullet(doc, lead, rest):
    p = doc.add_paragraph()
    pfmt(p, align="both", before=2, after=6, line=16)
    p.paragraph_format.left_indent = Cm(0.75)
    r = p.add_run("• " + lead)
    set_run(r, size=12, bold=True)
    r2 = p.add_run(" " + rest)
    set_run(r2, size=12)


def add_caption(doc, text):
    p = doc.add_paragraph()
    pfmt(p, align="center", before=4, after=10, line=14)
    r = p.add_run(text)
    set_run(r, size=11, bold=True)


def shade_cell(cell, fill=None):
    """No fill — college tables print with a clear/white background."""
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tcPr.append(shd)


def clear_table_banding(tbl):
    tblPr = tbl._tbl.tblPr
    for child in list(tblPr):
        if child.tag == qn("w:tblLook"):
            tblPr.remove(child)
    look = OxmlElement("w:tblLook")
    look.set(qn("w:firstRow"), "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "1")
    look.set(qn("w:noVBand"), "1")
    tblPr.append(look)


def set_cell(cell, text, *, header=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    pfmt(p, align="left", before=2, after=2, line=12)
    r = p.add_run(text)
    set_run(r, size=size, bold=header, color=RGBColor(0, 0, 0))


def table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    for child in list(tblPr):
        if child.tag in (qn("w:tblStyle"), qn("w:tblBorders")):
            tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(tbl)
    clear_table_banding(tbl)
    for i, h in enumerate(headers):
        shade_cell(tbl.rows[0].cells[i])
        set_cell(tbl.rows[0].cells[i], h, header=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            shade_cell(tbl.rows[ri + 1].cells[ci])
            set_cell(tbl.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    pfmt(hp, align="center", before=0, after=0, line=12)
    r = hp.add_run(TITLE)
    set_run(r, size=10)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    pfmt(fp, align="center", before=0, after=0, line=12)
    add_page_field(fp)
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "80")
    sectPr.append(pg)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = TNR
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    setup_header_footer(doc)

    add_title(doc, "CHAPTER 7", size=16, after=4, before=6)
    add_title(doc, "RESULTS, DISCUSSION, CONCLUSION AND FUTURE SCOPE", size=14, after=16)

    add_h(doc, "7.1 RESULTS AND DISCUSSION")
    add_body(
        doc,
        "The overarching objective of this project was to eliminate the reliance on subjective "
        "manual visual judgment by providing a continuous, objective risk assessment of bus stop "
        "overcrowding. The development, testing, and deployment phases—detailed across Chapters "
        "4, 5, and 6—confirmed that the integrated IoT and machine learning architecture "
        "successfully achieved this goal. Synthesizing the technical findings with the community "
        "impact assessment reveals a highly reliable, low-cost system that successfully bridges "
        "the gap between edge-level physical sensing and cloud-based predictive analytics.",
    )
    add_body(
        doc,
        "From a technical hardware perspective, the system demonstrated exceptional stability. "
        "Testing results from Chapter 5 yielded a 100% pass rate across all 11 unit tests "
        "(UT-01 to UT-11) and all 6 integration tests (IT-01 to IT-06). The ESP32 edge node "
        "effectively managed the physical sensing layer, confirming passenger crossings with a "
        "highly responsive 150 ms polling interval. By utilizing a strict logic filter that "
        "required two consecutive readings below 400 mm, the hardware successfully rejected "
        "transient sensor noise and prevented false crossing registrations. During early "
        "development, three critical bugs were systematically identified and resolved: a VL53L0X "
        "library API mismatch, a VN versus 3.3 V power rail miswiring that left sensors "
        "unpowered, and a V2/V3 virtual pin ownership conflict. Resolving these resulted in a "
        "finalized 3.3 V architecture and strict pin-ownership protocol that performed "
        "flawlessly during continuous deployment.",
    )
    add_body(
        doc,
        "The performance of the backend processing and communication layers was equally robust. "
        "The Python backend (brain.py) maintained a strict 5-second processing cycle, ensuring "
        "the system remained highly responsive to real-time microclimate and crowding changes. "
        "The communication bridge via Blynk Cloud recorded an average round trip latency of "
        "891 ms, well within acceptable operational limits. Crucially, the Telegram alerts were "
        "delivered to the transit authority in just 0.81 seconds. To ensure the system preserved "
        "a persistent historical record for future analysis and retraining, operational data was "
        "appended to the local bus_stop_log.csv file in effectively 0.00 ms per cycle.",
    )
    add_body(
        doc,
        "The most significant technical finding pertained to the machine learning model's "
        "performance and its precise alignment with the foundational risk formula. The "
        "RandomForestRegressor, configured with 100 trees and a random_state=42, executed its "
        "risk computation in just 10.51 ms, making it exceptionally well-suited for a live IoT "
        "pipeline. Because the model was trained on a dense, formula-generated synthetic grid of "
        "43,632 samples (covering 24 hours, 101 passenger counts, 2 rain states, and 9 "
        "temperatures), it perfectly approximated the non-linear environmental penalties of the "
        "system without overfitting. The model output exactly matched the expected formula "
        "results under the tested scenarios: predicting a 90.0% CRITICAL risk at full capacity "
        "(45 passengers) with rain, a 65.0% WARNING risk for 30 passengers with rain, and a "
        "baseline 15.0% NORMAL risk for 0 passengers with rain.",
    )
    add_body(
        doc,
        "From an operational and community perspective, the system performed exactly as intended "
        "during the simulated pilot deployment at the Dindoshi Bus Depot (Route 326) context. "
        "The strict boundary logic—where a score of exactly 75 remains classified as WARNING, "
        "and only a score above 75 triggers CRITICAL—was pivotal in preventing false-positive "
        "emergencies. Overcrowding alone (reaching 75) was intentionally not treated as an "
        "emergency; environmental pressure (rain, heat ≥33°C, or rush hour) was required to "
        "cross the critical threshold. When this threshold was breached, the community impact "
        "was immediate: on-site commuters were informed via the red LED indicator (with public "
        "audio firmware-disabled for safety), while remote transit authorities received automated "
        "Telegram alerts. The enforcement of a 60-second alert cooldown successfully prevented "
        "notification fatigue for the operators.",
    )
    add_body(
        doc,
        'Finally, the V7 "Bus Left (-30)" push button and the V8 "minutes since last crossing" '
        "display proved to be a highly effective human-in-the-loop design decision. Instead of "
        "utilizing an automatic count decay—which risks erasing waiting commuters if a bus "
        "reaches capacity and leaves stragglers behind—the system placed the final subtractive "
        "control in the hands of the human operator. By referencing the V8 inactivity timer as "
        "concrete evidence before pressing V7, transit staff maintained safe, contextual control "
        "over the system's tracking accuracy.",
    )

    add_h(doc, "7.2 ACHIEVEMENT OF OBJECTIVES")
    add_body(
        doc,
        "The project's original objectives, defined in Chapter 1, were systematically evaluated "
        "against the final built prototype to determine the extent of their completion. The "
        "results confirm that all primary objectives were successfully fulfilled.",
    )
    add_caption(doc, "Table 7.1: Status of Project Objectives")
    add_table(
        doc,
        ["Objective", "Status", "Evidence"],
        [
            [
                "1. To design a low-cost IoT hardware node for counting passengers and monitoring ambient temperature at a single bus stop.",
                "Met",
                "The ESP32 successfully integrates the VL53L0X (passenger crossings via GPIO21/22) and DHT22 (temperature via GPIO4) on a unified 3.3 V power rail.",
            ],
            [
                "2. To integrate live environmental context (weather and time) with local sensor data.",
                "Met",
                "The Python backend successfully fetches live Dindoshi weather via the OpenWeatherMap API, mapping WMO codes to a binary rain state.",
            ],
            [
                "3. To develop a Machine Learning model capable of predicting a continuous crowding risk score.",
                "Met",
                "A RandomForestRegressor trained on 43,632 samples accurately outputs a 0–98% continuous risk score matching the raw risk formula.",
            ],
            [
                "4. To provide transparent, real-time alerts to both on-site commuters and remote transit authorities.",
                "Met",
                "Local LEDs (GPIO25/26) provide on-site awareness, while Telegram Bot API alerts trigger perfectly under the above 75 strict CRITICAL rule.",
            ],
            [
                '5. To allow safe operator correction of passenger counts following a bus departure.',
                "Met",
                'The V7 "Bus Left" push button and V8 "minutes since last crossing" evidence gauge were successfully implemented on the Blynk dashboard.',
            ],
        ],
    )

    add_h(doc, "7.3 PROJECT CONTRIBUTIONS")
    add_body(
        doc,
        "This project introduces a highly accessible, stop-level predictive tool that fills a "
        "significant technical and operational gap in current public transit management. "
        "Previously, transit authorities relied on expensive, vehicle-mounted Industrial "
        "Automatic Passenger Counting (APC) systems, which measure fleet load but ignore the "
        "physical conditions at the shelter. Alternatively, commuters utilized transit "
        "applications like Chalo or Google Maps, which track vehicle GPS locations but remain "
        "entirely blind to the actual crowding pressure at a specific stop.",
    )
    add_body(
        doc,
        "The primary technical contribution of this project is the successful fusion of "
        "anonymous, edge-level Time-of-Flight (ToF) sensing with live cloud-based microclimate "
        "data and a Random Forest machine learning backend, all achieved using exceptionally "
        "low-cost hardware. By translating the feature set [hour_of_day, is_raining, "
        "passenger_count, temperature] into a transparent, actionable risk percentage, this "
        "approach replaces subjective visual judgment with objective, quantifiable data.",
    )
    add_body(
        doc,
        "The practical contribution to the community is the establishment of a proactive safety "
        "measure. Commuters no longer have to guess the severity of the crowding, and operators "
        "no longer have to wait for visual confirmation or commuter complaints to dispatch "
        "feeder buses. Beyond the pilot users at the Dindoshi Bus Depot, this architecture holds "
        "broader relevance for municipal bodies, such as the Mumbai Metropolitan Region "
        "Development Authority (MMRDA), and suburban transit networks. It provides a scalable "
        "blueprint for monitoring unstaffed, high-density pedestrian infrastructure without "
        "investing in invasive CCTV cameras or expensive fleet-wide technology upgrades.",
    )

    add_h(doc, "7.4 LIMITATIONS")
    add_body(
        doc,
        "While the prototype successfully meets its core requirements and functional "
        "specifications, several practical and technical limitations were identified during the "
        "development, testing, and deployment phases.",
    )
    add_bullet(
        doc,
        "Single Sensor Ghost Counting:",
        "Because the system utilizes a single VL53L0X Time-of-Flight sensor, it can only detect "
        "proximity crossings; it cannot explicitly calculate entry versus exit direction. An "
        "individual lingering or pacing within the sensor's 400 mm field of view could trigger "
        "ghost counts, artificially inflating the passenger tally. This hardware constraint "
        "makes the V7 operator correction essential for maintaining long-term accuracy.",
    )
    add_bullet(
        doc,
        "Estimated Capacity Constraint:",
        "The maximum capacity baseline of 45 passengers is a reasoned design estimate, not a "
        "physically measured metric for the specific Dindoshi shelter. It was calculated by "
        "estimating a typical shelter footprint of 36 m² and dividing it by 0.80 m² per "
        "passenger, which strictly aligns with the Level of Service C (LOS C) comfortable "
        "standing density standard defined by the transit-design Highway Capacity Manual.",
    )
    add_bullet(
        doc,
        "Synthetic Training Data:",
        "The Random Forest model was trained on a 43,632-sample formula-generated grid rather "
        "than historical, real-world crowding data. Because localized sensor-fusion data for "
        "this specific stop did not previously exist, the model currently acts as a highly "
        "accurate formula approximator rather than an engine discovering hidden real-world "
        "patterns.",
    )
    add_bullet(
        doc,
        "Wi-Fi Dependency:",
        "The backend predictive logic and alerting mechanisms depend entirely on continuous "
        "Wi-Fi connectivity. If the internet connection drops, the ESP32 edge node will "
        "successfully continue counting passengers locally, but the live weather fetching, risk "
        "computation, Telegram alerts, and Blynk dashboard updates will fail until the network "
        "connection is restored.",
    )

    add_h(doc, "7.5 FUTURE ENHANCEMENTS")
    add_body(
        doc,
        "The current prototype establishes a robust, foundational architecture that can be "
        "significantly expanded and improved through several targeted future enhancements.",
    )
    add_bullet(
        doc,
        "Directional Sensing Integration:",
        "Upgrading the hardware edge layer to utilize dual VL53L0X sensors, or transitioning to "
        "a multi-zone Time-of-Flight sensor such as the VL53L5CX, would allow the "
        "microcontroller to calculate entry and exit vectors. This would effectively solve the "
        "ghost-counting limitation, creating an autonomous, self-correcting tally that reduces "
        "the reliance on manual operator corrections.",
    )
    add_bullet(
        doc,
        "Real-World Model Retraining:",
        "The system currently appends the results of every 5-second processing cycle to a local "
        "bus_stop_log.csv file. Over time, this will generate a massive dataset of authentic "
        "commuter behavior at the Dindoshi stop. In the future, this accumulated CSV data can "
        "be used to retrain the Random Forest model, shifting it from a synthetic "
        "formula-approximation tool to an empirical, data-driven predictive model based on true "
        "historical trends.",
    )
    add_bullet(
        doc,
        "Cloud Database Migration:",
        "Transitioning the persistent storage architecture from a localized CSV file on a "
        "single backend machine to a managed cloud database (such as MySQL or Firebase) would "
        "allow for complex historical querying. This migration is a necessary step for "
        "supporting a scalable architecture that spans multiple bus stops and transit routes "
        "across the city.",
    )
    add_bullet(
        doc,
        "Solar Power Implementation:",
        "Replacing the current USB power dependency with a localized solar panel and battery "
        "management system would allow the ESP32 edge node to be deployed entirely off-grid, "
        "expanding the system's viability to remote or underdeveloped transit shelters.",
    )

    add_h(doc, "7.6 CONCLUSION")
    add_body(
        doc,
        "This project set out to resolve the lack of objective, real-time overcrowding awareness "
        "at public bus stops by replacing reactive, manual visual judgment with an automated "
        "predictive system. By integrating a low-cost ESP32 controller, an anonymous VL53L0X "
        "Time-of-Flight sensor, and a DHT22 module, the hardware successfully captured localized "
        "physical conditions at the edge. When fused with live OpenWeatherMap data and processed "
        "through a RandomForestRegressor machine learning model, the system generated highly "
        "accurate, continuous risk scores that dynamically accounted for rain, heat, and "
        "rush-hour pressures.",
    )
    add_body(
        doc,
        "The technical execution proved highly reliable, characterized by 100% test pass rates, "
        "rapid 150 ms sensor polling, and 10.51 ms predictive computations. The deployment "
        "impact on the community was immediate and highly practical: commuters gained on-site "
        "situational awareness via glanceable LED indicators, while transit authorities received "
        "automated Telegram alerts the moment the risk crossed the strict above 75 CRITICAL "
        "threshold. Ultimately, this IoT-based solution proves that scalable, low-cost sensor "
        "fusion can proactively safeguard commuter transit networks, providing vital "
        "overcrowding intelligence without compromising public privacy or requiring massive "
        "infrastructural investments in CCTV or industrial APC systems.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
