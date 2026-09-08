#!/usr/bin/env python3
"""Cover, certificate, proposal, TOC — friend's NMFC black-book front matter."""
from html import escape
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor, Twips
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

ROOT = Path("/workspace/cep-iot-bus-stop")
FIG = ROOT / "figures/front_matter"
CH = ROOT / "chapters"
LOGO_COVER = FIG / "logo_p1_0_129x211.png"
LOGO_HEAD = FIG / "logo_p2_0_102x177.png"
LOGO_WORD = ROOT / "figures/questionnaires/nmfc_logo.png"

RED = HexColor("#C00000")
BLUE = HexColor("#2E5A9C")
ORANGE = HexColor("#E36C09")
TNR = "Times-Roman"
TNRB = "Times-Bold"
TNRU = "Times-Bold"

TITLE = (
    "IoT-Based Bus Stop System for Public Transport Overcrowding "
    "Prediction Using Machine Learning"
)
STUDENTS = "Mr. Om Keluskar and Mr. Jashith Agre"
SEAT_A, SEAT_B = "24TIT078", "24TIT011"
SEATS = f"{SEAT_A} and {SEAT_B}"
GUIDE = "Prof. George Thekkevilayil"
GUIDE_COVER = "Mr. George Thekkevilayil"
COORD = "Prof. Vaishali Mishra"
YEAR = "2026–2027"
YEAR_SHORT = "2026-27"

TOC = [
    (
        "CHAPTER 1: INTRODUCTION AND COMMUNITY CONTEXT",
        1,
        [
            ("1.1. Project Background", 1),
            ("1.2. Community Problem Identification", 1),
            ("1.3. Community Need Assessment", 2),
            ("1.4. Existing System Analysis", 3),
            ("1.5. Problem Statement", 3),
            ("1.6. Proposed Solution", 3),
            ("1.7. Aim and Objectives", 4),
            ("1.8. Scope of Project", 4),
            ("1.9. Limitations", 5),
            ("1.10. Major Modules", 5),
            ("1.11. Success Criteria", 6),
            ("1.12. Expected Outcomes", 6),
            ("1.13. Chapter Summary", 7),
        ],
    ),
    (
        "CHAPTER 2: SYSTEM ANALYSIS AND REQUIREMENT ANALYSIS",
        8,
        [
            ("2.1. Introduction", 8),
            ("2.2. Literature Review", 8),
            ("2.3. Comparative Analysis", 11),
            ("2.4. Research Gap Identification", 11),
            ("2.5. Requirement Gathering", 12),
            ("2.6. Stakeholder Analysis", 13),
            ("2.7. Functional Requirements", 14),
            ("2.8. Nonfunctional Requirements", 15),
            ("2.9. Feasibility Study", 16),
            ("2.10. Software Requirements", 16),
            ("2.11. Hardware Requirements", 17),
            ("2.12. Project Schedule", 17),
            ("2.13. Development Methodology", 18),
        ],
    ),
    (
        "CHAPTER 3: SYSTEM DESIGN",
        19,
        [
            ("3.1. System Architecture", 19),
            ("3.2. System Workflow", 20),
            ("3.3. Block Diagram", 21),
            ("3.4. Functional Flow Diagram", 22),
            ("3.5. Circuit Diagram", 23),
            ("3.6. Component Interconnection Diagram", 25),
            ("3.7. Hardware Module Design", 27),
            ("3.8. Data Flow & Processing Logic", 30),
            ("3.9. Database Design (if applicable)", 31),
            ("3.10. Dashboard/UI Design", 32),
            ("3.11. Technology Stack", 33),
            ("3.12. Safety & Security Considerations", 34),
            ("3.13. Chapter Summary", 35),
        ],
    ),
    (
        "CHAPTER 4: DEVELOPMENT AND IMPLEMENTATION",
        36,
        [
            ("4.1. Development Environment", 36),
            ("4.2. Implementation Approach", 36),
            ("4.3. Hardware Assembly", 38),
            ("4.4. Firmware Development", 42),
            ("4.5. Database/Cloud Integration", 44),
            ("4.6. Dashboard Development", 46),
            ("4.7. Module-wise Implementation", 48),
            ("4.8. Algorithms and Control Logic", 50),
            ("4.9. Safety Features Implemented", 52),
            ("4.10. Prototype Development", 53),
            ("4.11. Prototype Photos & Demonstration", 54),
        ],
    ),
    (
        "CHAPTER 5: TESTING, VALIDATION AND PERFORMANCE EVALUATION",
        58,
        [
            ("5.1. Testing Strategy", 58),
            ("5.2. Unit Testing", 58),
            ("5.3. Integration Testing", 60),
            ("5.4. System Testing", 62),
            ("5.5. User Acceptance Testing", 64),
            ("5.6. Performance Evaluation", 65),
            ("5.7. Security Testing", 67),
            ("5.8. Validation Metrics", 68),
            ("5.9. Test Cases", 69),
            ("5.10. Defects Identified and corrections", 70),
            ("5.11. Chapter Summary", 71),
        ],
    ),
    (
        "CHAPTER 6: COMMUNITY DEPLOYMENT AND IMPACT ASSESSMENT",
        72,
        [
            ("6.1. Beneficiary Profile", 72),
            ("6.2. Community Engagement Activities", 72),
            ("6.3. Pilot Deployment", 73),
            ("6.4. User Training", 74),
            ("6.5. Feedback Collection", 75),
            ("6.6. Feedback Analysis", 75),
            ("6.7. Impact Assessment", 77),
            ("6.8. Outcome Achievement Matrix", 78),
            ("6.9. Sustainability Plan", 79),
        ],
    ),
    (
        "CHAPTER 7: RESULTS, DISCUSSION, CONCLUSION AND FUTURE SCOPE",
        80,
        [
            ("7.1. Results and Discussion", 80),
            ("7.2. Achievement of Objectives", 81),
            ("7.3. Project Contributions", 82),
            ("7.4. Limitations", 83),
            ("7.5. Future Enhancements", 84),
            ("7.6. Conclusion", 85),
        ],
    ),
]

BACK = [
    ("REFERENCES", "86"),
    ("ANNEXURES", "89 onwards"),
]


def border(c, inset=22, pagesize=letter):
    W, H = pagesize
    c.setStrokeColor(black)
    c.setLineWidth(1.0)
    c.rect(inset, inset, W - 2 * inset, H - 2 * inset)


def wrap(c, text, font, size, maxw):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if stringWidth(trial, font, size) <= maxw:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def letterhead(c, y_top):
    """Friend's red/blue header with logo at left."""
    W, H = letter
    if LOGO_HEAD.exists():
        c.drawImage(
            str(LOGO_HEAD),
            40,
            y_top - 78,
            width=42,
            height=72,
            mask="auto",
            preserveAspectRatio=True,
        )
    c.setFillColor(RED)
    c.setFont(TNRB, 16)
    c.drawCentredString(W / 2, y_top - 8, "Nirmala Memorial Foundation College")
    c.drawCentredString(W / 2, y_top - 26, "of Commerce and Science")
    c.setFillColor(BLUE)
    c.setFont(TNRB, 12)
    c.drawCentredString(W / 2, y_top - 44, "(Autonomous)")
    c.setFont(TNR, 9)
    c.drawCentredString(
        W / 2,
        y_top - 58,
        "Re-Accredited by NAAC with B++ Grade & ISO 9001:2015 Certified",
    )
    c.drawCentredString(
        W / 2,
        y_top - 70,
        "Recognized under Section 2 (f) & 12 (B) of the UGC Act, 1956",
    )
    c.setStrokeColor(ORANGE)
    c.setLineWidth(1.1)
    c.line(40, y_top - 84, W - 40, y_top - 84)
    c.setFillColor(BLUE)
    c.setFont(TNR, 8.5)
    c.drawCentredString(
        W / 2,
        y_top - 98,
        "D S Road, Asha Nagar, Thakur Complex, Kandivali (East), "
        "Mumbai-400101  Tel: 022-69436400",
    )
    c.line(40, y_top - 108, W - 40, y_top - 108)
    c.setFillColor(black)
    return y_top - 124


def draw_cover(c):
    W, H = letter
    border(c)
    y = H - 56
    c.setFillColor(black)
    c.setFont(TNRB, 13)
    for line in wrap(
        c,
        "NIRMALA MEMORIAL FOUNDATION COLLEGE OF COMMERCE AND SCIENCE",
        TNRB,
        13,
        W - 80,
    ):
        c.drawCentredString(W / 2, y, line)
        y -= 16
    c.setFont(TNRB, 12)
    c.drawCentredString(W / 2, y, "(AUTONOMOUS)")
    y -= 16
    c.setFont(TNR, 9)
    for line in wrap(
        c,
        "Re-Accredited by NAAC with B++ Grade & ISO 9001:2015 Certified "
        "Recognized under Section 2 (f) & 12 (B) of the UGC Act, 1956.",
        TNR,
        9,
        W - 90,
    ):
        c.drawCentredString(W / 2, y, line)
        y -= 12
    c.setFont(TNR, 10)
    c.drawCentredString(W / 2, y, "Kandivali (East), Mumbai-400101.")
    y -= 28
    if LOGO_COVER.exists():
        c.drawImage(
            str(LOGO_COVER),
            W / 2 - 38,
            y - 118,
            width=76,
            height=124,
            mask="auto",
            preserveAspectRatio=True,
        )
        y -= 140
    c.setFont(TNRB, 16)
    for line in wrap(c, TITLE, TNRB, 16, W - 90):
        c.drawCentredString(W / 2, y, line)
        y -= 20
    y -= 10
    c.setFont(TNRB, 13)
    c.drawCentredString(W / 2, y, "A Project Report")
    y -= 18
    c.setFont(TNR, 11)
    for line in wrap(
        c,
        "Submitted in partial fulfillment of the requirements "
        "For the degree of B.Sc. Information Technology aligned with "
        "the NEP-2020 Community Engagement Project (CEP)",
        TNR,
        11,
        W - 100,
    ):
        c.drawCentredString(W / 2, y, line)
        y -= 15
    y -= 12
    c.setFont(TNR, 12)
    c.drawCentredString(W / 2, y, "By")
    y -= 18
    c.setFont(TNRB, 13)
    c.drawCentredString(W / 2, y, STUDENTS)
    y -= 16
    c.setFont(TNR, 12)
    c.drawCentredString(W / 2, y, f"Seat No: {SEAT_A}, {SEAT_B}")
    y -= 15
    c.drawCentredString(W / 2, y, "Semester: V")
    y -= 28
    c.drawCentredString(W / 2, y, "Under the esteemed guidance of")
    y -= 18
    c.setFont(TNRB, 13)
    c.drawCentredString(W / 2, y, GUIDE_COVER)
    y -= 16
    c.drawCentredString(W / 2, y, "Assistant Professor")
    y -= 36
    c.setFont(TNRB, 12)
    c.drawCentredString(W / 2, y, "BACHELOR OF SCIENCE INFORMATION TECHNOLOGY")
    y -= 18
    c.drawCentredString(W / 2, y, f"Academic Year: {YEAR}")


def draw_certificate(c):
    W, H = letter
    border(c, 20)
    y = letterhead(c, H - 42)
    y -= 18
    c.setFont(TNRB, 16)
    c.drawCentredString(W / 2, y, "CERTIFICATE OF PROJECT")
    y -= 18
    c.drawCentredString(W / 2, y, "COMPLETION")
    c.setStrokeColor(black)
    c.setLineWidth(0.8)
    tw = stringWidth("CERTIFICATE OF PROJECT COMPLETION", TNRB, 16)
    # underline the two-line title block
    c.line(W / 2 - 160, y - 4, W / 2 + 160, y - 4)
    y -= 22
    c.setFont(TNRB, 13)
    c.drawCentredString(W / 2, y, "B.Sc. (Information Technology)")
    c.line(W / 2 - 120, y - 3, W / 2 + 120, y - 3)
    y -= 36
    body = (
        f'This is to certify that {STUDENTS}, students of T.Y.B.Sc. '
        f"(Information Technology), Semester V, bearing Seat No. {SEAT_A} and "
        f'{SEAT_B}, has successfully completed the Project Work titled "{TITLE}." '
        f"The project has been submitted in partial fulfillment of the requirements "
        f"prescribed in the syllabus for the award of the Degree of Bachelor of "
        f"Science in Information Technology for the academic year {YEAR}. The "
        f"practical and project work has been carried out successfully under the "
        f"supervision of the undersigned."
    )
    c.setFont(TNR, 12)
    for line in wrap(c, body, TNR, 12, W - 90):
        c.drawString(48, y, line)
        y -= 16
    y = 175
    xs = [70, W / 2, W - 70]
    labels = ["Internal Examiner", "External Examiner", "Co-ordinator"]
    names = [GUIDE, "", COORD]
    c.setStrokeColor(black)
    c.setLineWidth(1)
    for x, lab, name in zip(xs, labels, names):
        c.setFont(TNRB, 11)
        c.drawCentredString(x, y + 14, lab)
        c.line(x - 70, y, x + 70, y)
        if name:
            c.setFont(TNRB, 11)
            c.drawCentredString(x, y - 16, name)
    c.setFont(TNR, 11)
    c.drawCentredString(W / 2, 48, "College seal and Date")


def draw_proposal(c):
    W, H = letter
    border(c, 20)
    y = letterhead(c, H - 42)
    y -= 16
    c.setFont(TNRB, 14)
    c.drawCentredString(W / 2, y, "PROJECT PROPOSAL APPROVAL SHEET")
    c.setStrokeColor(black)
    c.setLineWidth(0.8)
    c.line(W / 2 - 155, y - 3, W / 2 + 155, y - 3)
    y -= 20
    c.setFont(TNRB, 12)
    c.drawCentredString(W / 2, y, "Department of Information Technology")
    y -= 16
    c.drawCentredString(W / 2, y, f"Academic Year {YEAR_SHORT}")
    y -= 36
    rows = [
        (f"1.  Name of the Students:  {STUDENTS}", False),
        (f"2.  Seat Number:  {SEAT_A}, {SEAT_B}", False),
        ("3.  Is this your first submission?     Yes            No", True),
        (f"4.  Name of the Guide:  {GUIDE}", False),
        ("5.  Teaching Experience of the Guide:  26 Years", False),
    ]
    c.setFont(TNR, 12)
    for text, boxes in rows:
        c.setFont(TNRB, 12)
        c.drawString(50, y, text)
        if boxes:
            c.setLineWidth(1)
            c.rect(318, y - 1, 11, 11, stroke=1, fill=0)
            c.rect(400, y - 1, 11, 11, stroke=1, fill=0)
        y -= 22
    y -= 18
    para = (
        f'The proposed project titled "{TITLE}." submitted by the above students, '
        f"has been examined and found suitable for the partial fulfilment of the "
        f"requirements for Semester V of the B.Sc IT Degree Programme. The project "
        f"is hereby approved for implementation under the guidance of the undersigned."
    )
    c.setFont(TNR, 12)
    for line in wrap(c, para, TNR, 12, W - 100):
        c.drawString(50, y, line)
        y -= 16
    y = 150
    c.setStrokeColor(black)
    c.setLineWidth(1)
    c.line(50, y, 230, y)
    c.line(W - 230, y, W - 50, y)
    c.setFont(TNR, 11)
    c.drawString(50, y - 16, "Signature of Guide")
    c.drawRightString(W - 50, y - 16, "Signature of Co-ordinator")
    c.drawString(50, y - 36, "Date:")
    c.drawRightString(W - 50, y - 36, "Date:")


def _toc_dots(c, x0, x1, y):
    c.setFont(TNR, 10)
    x = x0
    while x < x1:
        c.drawString(x, y, ".")
        x += 3.8


def toc_entry(c, left, right, y, text, page, *, bold=False, size=12, section=False):
    """Dotted-leader TOC row. Chapter titles wrap; section numbers stay bold."""
    num = str(page)
    font = TNRB if bold else TNR
    nw = stringWidth(num, TNRB if bold else TNR, size)
    usable = right - left - nw - 14
    prefix = ""
    prefix_w = 0
    body = text
    if section:
        head, _, tail = text.partition(" ")
        prefix = head + " "
        prefix_w = stringWidth(prefix, TNRB, size)
        body = tail
        lines = wrap(c, body, TNR, size, max(40, usable - prefix_w))
    else:
        lines = wrap(c, text, font, size, usable)
    if not lines:
        lines = [""]
    for i, line in enumerate(lines):
        last = i == len(lines) - 1
        if section and i == 0:
            c.setFont(TNRB, size)
            c.drawString(left, y, prefix.rstrip())
            c.setFont(TNR, size)
            c.drawString(left + prefix_w, y, line)
            end_x = left + prefix_w + stringWidth(line, TNR, size)
        else:
            use_font = font if not section else TNR
            x = left + prefix_w if section else left
            c.setFont(use_font, size)
            c.drawString(x, y, line)
            end_x = x + stringWidth(line, use_font, size)
        if last:
            _toc_dots(c, end_x + 4, right - nw - 6, y)
            c.setFont(TNRB if bold else TNR, size)
            c.drawRightString(right, y, num)
        y -= 16 if section else 15
    return y


def draw_toc_pages(c):
    W, H = A4
    left, right = 48, W - 48
    first = True

    def new_page():
        nonlocal first
        if not first:
            c.showPage()
        first = False
        border(c, pagesize=A4)
        return H - 48

    y = new_page()
    c.setFont(TNRB, 16)
    c.drawCentredString(W / 2, y, "TABLE OF CONTENTS")
    y -= 30
    for title, page, secs in TOC:
        if title.startswith("CHAPTER 3") or title.startswith("CHAPTER 5"):
            y = new_page()
        y = toc_entry(c, left, right, y, title, page, bold=True, size=12)
        y -= 2
        for s, sp in secs:
            if y < 68:
                y = new_page()
            y = toc_entry(c, left + 18, right, y, s, sp, section=True, size=12)
        y -= 10
    y -= 4
    for title, page in BACK:
        y = toc_entry(c, left, right, y, title, page, bold=True, size=12)
        y -= 4


def build_pdfs():
    CH.mkdir(parents=True, exist_ok=True)
    cover = CH / "Cover_Page.pdf"
    c = canvas.Canvas(str(cover), pagesize=letter)
    draw_cover(c)
    c.save()
    cert = CH / "Certificate_of_Project_Completion.pdf"
    # Certificate, proposal, and questionnaires are copied from the
    # original college PDF by scripts/restore_original_forms.py.
    prop = CH / "Project_Proposal_Approval_Sheet.pdf"
    # Proposal/survey pages come from the original college PDFs
    # (scripts/restore_original_forms.py). Do not redraw them here.
    toc = CH / "Table_of_Contents.pdf"
    c = canvas.Canvas(str(toc), pagesize=A4)
    draw_toc_pages(c)
    c.save()
    export_toc_pngs()
    # combined
    import pymupdf
    out = CH / "Front_Matter.pdf"
    merged = pymupdf.open()
    for p in (cover, cert, prop, toc):
        merged.insert_pdf(pymupdf.open(p))
    # append existing survey + feedback
    for extra in (
        CH / "CEP_Survey_Questionnaire.pdf",
        CH / "CEP_Feedback_Questionnaire.pdf",
    ):
        if extra.exists():
            merged.insert_pdf(pymupdf.open(extra))
    merged.save(out)
    print("wrote", cover, cert, prop, toc, out)


def set_run(run, *, size=12, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def pfmt(p, align="center", before=0, after=6, line=16):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]


def page_border(section):
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    section._sectPr.append(borders)


def add_toc_para(doc, text, page, bold=False, indent=0, section=False):
    p = doc.add_paragraph()
    pfmt(p, align="left", before=0, after=2, line=18)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    tab = p.paragraph_format.tab_stops.add_tab_stop
    tab(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if section:
        head, _, tail = text.partition(" ")
        r = p.add_run(head + " ")
        set_run(r, size=12, bold=True)
        r = p.add_run(tail)
        set_run(r, size=12, bold=False)
        r2 = p.add_run("\t" + str(page))
        set_run(r2, size=12, bold=False)
    else:
        r = p.add_run(text)
        set_run(r, size=12, bold=bold)
        r2 = p.add_run("\t" + str(page))
        set_run(r2, size=12, bold=bold)
    return p


def write_proposal_word(doc):
    p = doc.add_paragraph()
    pfmt(p, after=4)
    set_run(p.add_run("PROJECT PROPOSAL APPROVAL SHEET"), size=14, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run("Department of Information Technology"), size=12, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=14)
    set_run(p.add_run(f"Academic Year {YEAR_SHORT}"), size=12, bold=True)
    for line in [
        f"1.  Name of the Students:  {STUDENTS}",
        f"2.  Seat Number:  {SEAT_A}, {SEAT_B}",
        "3.  Is this your first submission?     Yes  ☐     No  ☐",
        f"4.  Name of the Guide:  {GUIDE}",
        "5.  Teaching Experience of the Guide:  26 Years",
    ]:
        p = doc.add_paragraph()
        pfmt(p, align="left", after=6, line=16)
        set_run(p.add_run(line), size=12, bold=True)
    p = doc.add_paragraph()
    pfmt(p, align="both", before=10, after=12, line=18)
    set_run(
        p.add_run(
            f'The proposed project titled "{TITLE}." submitted by the above students, '
            f"has been examined and found suitable for the partial fulfilment of the "
            f"requirements for Semester V of the B.Sc IT Degree Programme. The project "
            f"is hereby approved for implementation under the guidance of the undersigned."
        ),
        size=12,
    )
    p = doc.add_paragraph()
    pfmt(p, align="left", after=4)
    set_run(p.add_run("Signature of Guide                          Signature of Co-ordinator"), size=11)
    p = doc.add_paragraph()
    pfmt(p, align="left")
    set_run(p.add_run("Date:                                                Date:"), size=11)


def build_word():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    page_border(sec)

    # COVER
    for line, size, bold in [
        ("NIRMALA MEMORIAL FOUNDATION COLLEGE OF", 13, True),
        ("COMMERCE AND SCIENCE", 13, True),
        ("(AUTONOMOUS)", 12, True),
    ]:
        p = doc.add_paragraph()
        pfmt(p, after=2, line=16)
        set_run(p.add_run(line), size=size, bold=bold)
    p = doc.add_paragraph()
    pfmt(p, after=2, line=12)
    set_run(
        p.add_run(
            "Re-Accredited by NAAC with B++ Grade & ISO 9001:2015 Certified "
            "Recognized under Section 2 (f) & 12 (B) of the UGC Act, 1956."
        ),
        size=9,
    )
    p = doc.add_paragraph()
    pfmt(p, after=10, line=14)
    set_run(p.add_run("Kandivali (East), Mumbai-400101."), size=10)
    if LOGO_COVER.exists():
        p = doc.add_paragraph()
        pfmt(p, after=10)
        p.add_run().add_picture(str(LOGO_COVER), height=Cm(3.8))
    p = doc.add_paragraph()
    pfmt(p, after=10, line=20)
    set_run(p.add_run(TITLE), size=16, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=4)
    set_run(p.add_run("A Project Report"), size=13, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=10, line=16)
    set_run(
        p.add_run(
            "Submitted in partial fulfillment of the requirements for the degree of "
            "B.Sc. Information Technology aligned with the NEP-2020 Community "
            "Engagement Project (CEP)"
        ),
        size=11,
    )
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run("By"), size=12)
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run(STUDENTS), size=13, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run(f"Seat No: {SEAT_A}, {SEAT_B}"), size=12)
    p = doc.add_paragraph()
    pfmt(p, after=10)
    set_run(p.add_run("Semester: V"), size=12)
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run("Under the esteemed guidance of"), size=12)
    p = doc.add_paragraph()
    pfmt(p, after=2)
    set_run(p.add_run(GUIDE_COVER), size=13, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=14)
    set_run(p.add_run("Assistant Professor"), size=13, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=4)
    set_run(p.add_run("BACHELOR OF SCIENCE INFORMATION TECHNOLOGY"), size=12, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=0)
    set_run(p.add_run(f"Academic Year: {YEAR}"), size=12, bold=True)

    doc.add_page_break()
    # CERT
    p = doc.add_paragraph()
    pfmt(p, after=2, line=18)
    set_run(p.add_run("CERTIFICATE OF PROJECT COMPLETION"), size=16, bold=True)
    p = doc.add_paragraph()
    pfmt(p, after=12)
    set_run(p.add_run("B.Sc. (Information Technology)"), size=13, bold=True)
    p = doc.add_paragraph()
    pfmt(p, align="both", after=20, line=18)
    set_run(
        p.add_run(
            f"This is to certify that {STUDENTS}, students of T.Y.B.Sc. "
            f"(Information Technology), Semester V, bearing Seat No. {SEAT_A} and "
            f'{SEAT_B}, has successfully completed the Project Work titled "{TITLE}." '
            f"The project has been submitted in partial fulfillment of the requirements "
            f"prescribed in the syllabus for the award of the Degree of Bachelor of "
            f"Science in Information Technology for the academic year {YEAR}. The "
            f"practical and project work has been carried out successfully under the "
            f"supervision of the undersigned."
        ),
        size=12,
    )
    p = doc.add_paragraph()
    pfmt(p, after=20, line=16)
    set_run(
        p.add_run(
            f"Internal Examiner                    External Examiner                    Co-ordinator\n"
            f"{GUIDE}                                                                      {COORD}"
        ),
        size=11,
        bold=True,
    )
    p = doc.add_paragraph()
    pfmt(p)
    set_run(p.add_run("College seal and Date"), size=11)

    doc.add_page_break()
    write_proposal_word(doc)

    doc.add_page_break()
    write_toc_body(doc)

    out = CH / "Front_Matter.docx"
    doc.save(out)
    print("wrote", out)


def write_toc_body(doc, page_breaks=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1.0
    _set_run_simple(p.add_run("TABLE OF CONTENTS"), size=16, bold=True)
    for title, page, secs in TOC:
        if page_breaks and (
            title.startswith("CHAPTER 3") or title.startswith("CHAPTER 5")
        ):
            doc.add_page_break()
        add_leader_para(doc, title, page, bold=True)
        for s, sp in secs:
            add_leader_para(doc, s, sp, section=True)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
    for title, page in BACK:
        add_leader_para(doc, title, page, bold=True)


def _set_run_simple(run, *, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def _leader_count(title, page, *, size=12, bold=False, indent_cm=0.0):
    font = TNRB if bold else TNR
    target = 490
    indent_pts = indent_cm * 28.35
    left = f"{title} "
    right = str(page)
    avail = target - indent_pts - stringWidth(left, font, size) - stringWidth(
        right, font, size
    )
    unit = stringWidth(". ", TNR, size)
    return max(8, int(avail / unit))


def add_leader_para(doc, text, page, *, bold=False, section=False):
    """Dotted leaders as real '. .' characters so Google Docs keeps them."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.08
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    indent = 0.7 if section else 0.0
    if indent:
        pf.left_indent = Cm(indent)
    dots = ". " * _leader_count(
        text, page, size=12, bold=bold and not section, indent_cm=indent
    )
    page_s = str(page)
    if section:
        head, _, tail = text.partition(" ")
        r = p.add_run(head + " ")
        _set_run_simple(r, size=12, bold=True)
        r = p.add_run(f"{tail} {dots}{page_s}")
        _set_run_simple(r, size=12, bold=False)
    else:
        r = p.add_run(f"{text} {dots}{page_s}")
        _set_run_simple(r, size=12, bold=bold)
    return p


DOTS = ". " * 16


def build_toc_word():
    """Editable TOC: title | typed dots | page. Change font from the toolbar."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(20)
    sec.right_margin = Mm(20)
    sec.top_margin = Mm(18)
    sec.bottom_margin = Mm(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    _set_run_simple(title.add_run("TABLE OF CONTENTS"), size=16, bold=True)

    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint.paragraph_format.space_after = Pt(10)
    r = hint.add_run(
        "Select all (Ctrl+A) then change Font / Font size in the toolbar. "
        "Edit any title or page number. Add or delete dots in the middle column."
    )
    _set_run_simple(r, size=10, bold=False)
    r.italic = True

    table = doc.add_table(rows=0, cols=3)
    table.allow_autofit = True

    def add_row(left, page, *, bold=False, section=False):
        row = table.add_row()
        c0, c1, c2 = row.cells
        c0.width = Mm(95)
        c1.width = Mm(55)
        c2.width = Mm(18)
        p = c0.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if section:
            p.paragraph_format.left_indent = Cm(0.5)
            head, _, tail = left.partition(" ")
            r = p.add_run(head + " ")
            _set_run_simple(r, size=12, bold=True)
            r = p.add_run(tail)
            _set_run_simple(r, size=12, bold=False)
        else:
            _set_run_simple(p.add_run(left), size=12, bold=bold)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        _set_run_simple(p1.add_run(DOTS), size=12, bold=False)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        _set_run_simple(p2.add_run(str(page)), size=12, bold=bold)

    for ch_title, ch_page, secs in TOC:
        add_row(ch_title, ch_page, bold=True)
        for s, sp in secs:
            add_row(s, sp, section=True)
    for name, page in BACK:
        add_row(name, page, bold=True)

    out = CH / "Table_of_Contents.docx"
    doc.save(out)
    print("wrote", out)
    build_toc_html()
    import shutil

    art = Path("/opt/cursor/artifacts")
    art.mkdir(parents=True, exist_ok=True)
    shutil.copy(out, art / "Table_of_Contents_EDITABLE.docx")


def build_toc_html():
    def row(left, page, *, chapter=False, section=False):
        pad = "padding-left:18px;" if section else ""
        weight = "font-weight:bold;" if chapter else "font-weight:normal;"
        if section:
            head, _, tail = left.partition(" ")
            title = f"<b>{escape(head)}</b> {escape(tail)}"
        else:
            title = escape(left)
        return (
            "<tr>"
            f'<td style="{pad}{weight}padding:3px 6px 3px 0;">{title}</td>'
            f'<td style="padding:3px;">{DOTS}</td>'
            f'<td style="text-align:right;{weight}padding:3px 0;width:48px;">'
            f"{escape(str(page))}</td>"
            "</tr>"
        )

    parts = [
        "<!DOCTYPE html>",
        '<html lang="en"><head><meta charset="utf-8"/>',
        "<title>TABLE OF CONTENTS</title></head><body>",
        '<p style="text-align:center;font-family:\'Times New Roman\',Times,serif;'
        'font-size:16pt;font-weight:bold;">TABLE OF CONTENTS</p>',
        '<p style="text-align:center;font-family:\'Times New Roman\',Times,serif;'
        'font-size:10pt;font-style:italic;">'
        "Open with Google Docs. Click any word to edit. Select all, then change "
        "font size in the toolbar. Dots are normal text in the middle column."
        "</p>",
        '<table style="width:100%;border-collapse:collapse;'
        "font-family:'Times New Roman',Times,serif;font-size:12pt;\">",
    ]
    for ch_title, ch_page, secs in TOC:
        parts.append(row(ch_title, ch_page, chapter=True))
        for s, sp in secs:
            parts.append(row(s, sp, section=True))
    for name, page in BACK:
        parts.append(row(name, page, chapter=True))
    parts.append("</table></body></html>")
    out = CH / "Table_of_Contents.html"
    out.write_text("\n".join(parts), encoding="utf-8")
    print("wrote", out)
    import shutil

    shutil.copy(out, Path("/opt/cursor/artifacts") / "Table_of_Contents.html")


def export_toc_pngs():
    import pymupdf
    import zipfile
    import shutil

    pdf = pymupdf.open(CH / "Table_of_Contents.pdf")
    img_dir = CH / "toc_pages"
    img_dir.mkdir(parents=True, exist_ok=True)
    art = Path("/opt/cursor/artifacts")
    art.mkdir(parents=True, exist_ok=True)
    shots = Path("/opt/cursor/artifacts/screenshots")
    shots.mkdir(parents=True, exist_ok=True)
    pngs = []
    for i, page in enumerate(pdf):
        pix = page.get_pixmap(matrix=pymupdf.Matrix(2.4, 2.4), alpha=False)
        name = f"TOC_page_{i + 1}.png"
        dest = img_dir / name
        pix.save(str(dest))
        pngs.append(dest)
        shutil.copy(dest, art / name)
        pix2 = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False)
        pix2.save(str(shots / name))
    shutil.copy(CH / "Table_of_Contents.pdf", art / "Table_of_Contents.pdf")
    zpath = art / "TOC_insert_into_Google_Docs.zip"
    with zipfile.ZipFile(zpath, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(CH / "Table_of_Contents.pdf", "Table_of_Contents.pdf")
        for p in pngs:
            zf.write(p, p.name)
    print("wrote", img_dir, zpath)
    return pngs


def build_toc_pages_docx():
    """One A4 Word page per TOC PNG — Google Docs must not reflow text."""
    pngs = list((CH / "toc_pages").glob("TOC_page_*.png"))
    pngs.sort()
    if not pngs:
        pngs = export_toc_pngs()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(0)
    sec.right_margin = Mm(0)
    sec.top_margin = Mm(0)
    sec.bottom_margin = Mm(0)
    for i, img in enumerate(pngs):
        if i:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        p.add_run().add_picture(str(img), width=Mm(210), height=Mm(297))
    out = CH / "Table_of_Contents_PAGES.docx"
    doc.save(out)
    Path("/opt/cursor/artifacts").mkdir(parents=True, exist_ok=True)
    import shutil

    shutil.copy(out, Path("/opt/cursor/artifacts") / out.name)
    print("wrote", out)


if __name__ == "__main__":
    build_pdfs()
    build_word()
    build_toc_word()
    build_toc_pages_docx()
