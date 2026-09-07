#!/usr/bin/env python3
"""College CEP survey + feedback questionnaire pages (friend's layout)."""
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

ROOT = Path("/workspace/cep-iot-bus-stop")
LOGO_SRC = ROOT / "figures/questionnaires/nmfc_logo.png"
LOGO_WHITE = ROOT / "figures/questionnaires/nmfc_logo_white.png"
OUT_DIR = ROOT / "chapters"
SURVEY_PDF = OUT_DIR / "CEP_Survey_Questionnaire.pdf"
FEEDBACK_PDF = OUT_DIR / "CEP_Feedback_Questionnaire.pdf"
SURVEY_DOCX = OUT_DIR / "CEP_Survey_Questionnaire.docx"
FEEDBACK_DOCX = OUT_DIR / "CEP_Feedback_Questionnaire.docx"

NAVY = HexColor("#1F4E79")
TNR = "Times-Roman"
TNRB = "Times-Bold"

STUDENTS = "Mr. Om Keluskar, Mr. Jashith Agre"
SEATS = "24TIT078, 24TIT011"

SURVEY = [
    (
        "How frequently do you experience overcrowding at Dindoshi Bus Depot or the BEST Route 326 bus stop?",
        ["Every Day", "Several Times a Week", "Occasionally", "Rarely"],
    ),
    (
        "Which mode of transport do you use most frequently to reach this stop?",
        ["BEST Bus (Route 326)", "Other BEST Bus", "Train then BEST Bus", "Walking"],
    ),
    (
        "What is the biggest problem you face because of overcrowding at the bus stop?",
        [
            "Missed bus / cannot board",
            "Pushing or injury risk",
            "Long uncomfortable wait",
            "No crowding warning",
        ],
    ),
    (
        "Have you ever been unable to board a bus because the stop or doorway was too crowded?",
        ["Yes, Multiple Times", "Yes, Once", "No", "Not Applicable"],
    ),
    (
        "How safe do you feel while waiting at an overcrowded bus stop?",
        ["Very Safe", "Safe", "Unsafe", "Very Unsafe"],
    ),
    (
        "Have you ever reported overcrowding to depot staff or BEST officials?",
        ["Yes", "No", "Didn't Know How", "Never Felt It Would Help"],
    ),
    (
        "If you did report overcrowding, how satisfied were you with the response?",
        ["Very Satisfied", "Satisfied", "Dissatisfied", "Never Reported"],
    ),
    (
        "How important is it to warn commuters as soon as overcrowding is detected?",
        ["Extremely Important", "Important", "Slightly Important", "Not Important"],
    ),
    (
        "Do you think rain, heat, or peak hours make overcrowding at the stop more dangerous?",
        ["Strongly Agree", "Agree", "Disagree", "Strongly Disagree"],
    ),
    (
        "Would you trust an automated IoT system that predicts overcrowding using on-site lights, without requiring a mobile app?",
        ["Definitely Yes", "Probably Yes", "Probably No", "Definitely No"],
    ),
    (
        "Which feature would be most useful in a smart bus-stop overcrowding system?",
        [
            "On-site LED display",
            "Telegram staff alerts",
            "Live risk percentage",
            "Bus-left operator button",
        ],
    ),
    (
        "Do you believe an IoT-based overcrowding prediction system can improve commuter safety at BEST bus stops?",
        ["Strongly Agree", "Agree", "Disagree", "Strongly Disagree"],
    ),
]

FEEDBACK = [
    (
        "How easy was it to understand the purpose of the IoT bus-stop overcrowding system?",
        ["Very Easy", "Easy", "Neutral", "Difficult"],
    ),
    (
        "How useful do you think automatic passenger counting using the IoT sensor would be for overcrowding awareness?",
        ["Very Useful", "Useful", "Moderately Useful", "Not Useful at All"],
    ),
    (
        "How clear and understandable was the demonstrated overcrowding-prediction process?",
        ["Very Clear", "Clear", "Neutral", "Unclear"],
    ),
    (
        "How useful were the on-site red and green LED indicators for understanding crowding?",
        ["Very Useful", "Useful", "Moderately Useful", "Not Useful at All"],
    ),
    (
        "How useful would real-time Telegram alerts be for depot staff when crowding becomes CRITICAL?",
        ["Very Useful", "Useful", "Moderately Useful", "Not Useful at All"],
    ),
    (
        "How easy was it to understand the crowding information shown on the dashboard (count, weather, risk, status)?",
        ["Very Easy", "Easy", "Neutral", "Difficult"],
    ),
    (
        "How clearly did the hardware demonstration show how the sensor detects passenger crossings?",
        ["Very Clear", "Clear", "Neutral", "Unclear"],
    ),
    (
        "How confident are you that this system can warn of overcrowding faster than manual visual guessing?",
        ["Very Confident", "Confident", "Moderately Confident", "Slightly Confident"],
    ),
    (
        "Would you be willing to use this system if it remained available at Dindoshi Route 326?",
        ["Definitely Yes", "Probably Yes", "Not Sure", "Definitely No"],
    ),
    (
        "How clear and easy to understand were the NORMAL, WARNING, and CRITICAL status lights?",
        ["Very Clear", "Clear", "Neutral", "Unclear"],
    ),
    (
        "How useful do you think this system would be for improving crowding awareness at your bus stop?",
        ["Very Useful", "Useful", "Moderately Useful", "Not Useful at All"],
    ),
    (
        "Overall, how satisfied are you with the IoT bus-stop overcrowding system and its potential for commuter safety?",
        ["Very Satisfied", "Satisfied", "Neutral", "Dissatisfied"],
    ),
]


def prepare_logo():
    im = Image.open(LOGO_SRC).convert("RGBA")
    px = im.load()
    w, h = im.size
    for y in range(h):
        for x in range(w):
            r, g, b, a = px[x, y]
            if r < 40 and g < 40 and b < 40:
                px[x, y] = (255, 255, 255, 0)
    im.save(LOGO_WHITE)
    return LOGO_WHITE


def wrap(c, text, font, size, max_width):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if c.stringWidth(trial, font, size) <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [text]


def draw_page(c, *, title, questions, page_no, logo_path):
    W, H = letter
    margin = 0.55 * inch
    c.setStrokeColor(black)
    c.setLineWidth(1.1)
    c.rect(margin, margin, W - 2 * margin, H - 2 * margin)

    logo_h = 0.72 * inch
    logo_w = logo_h * (696 / 300) * 0.85
    c.drawImage(
        str(logo_path),
        margin + 8,
        H - margin - logo_h - 6,
        width=logo_w,
        height=logo_h,
        mask="auto",
        preserveAspectRatio=True,
    )

    hx = margin + logo_w + 14
    c.setFillColor(NAVY)
    c.setFont(TNRB, 13)
    c.drawString(hx, H - margin - 22, "Nirmala Memorial Foundation College of")
    c.drawString(hx, H - margin - 38, "Commerce and Science (Autonomous)")
    c.setFillColor(black)
    c.setFont(TNR, 10)
    c.drawString(hx, H - margin - 54, "Department of Information Technology")
    c.setFont(TNRB, 10)
    c.drawString(hx, H - margin - 68, "Community Engagement Project Questionnaire")

    box_top = H - margin - 86
    box_h = 52
    c.setLineWidth(0.8)
    c.rect(margin + 8, box_top - box_h, W - 2 * margin - 16, box_h)
    mid = W / 2
    c.line(mid, box_top - box_h, mid, box_top)
    c.setFont(TNRB, 9)
    c.drawString(margin + 14, box_top - 14, "Name of Student(s):")
    c.setFont(TNR, 9)
    c.drawString(margin + 118, box_top - 14, STUDENTS)
    c.setFont(TNRB, 9)
    c.drawString(margin + 14, box_top - 30, "Seat No:")
    c.setFont(TNR, 9)
    c.drawString(margin + 62, box_top - 30, SEATS)
    c.setFont(TNRB, 9)
    c.drawString(mid + 8, box_top - 14, "Name of Participant:")
    c.setStrokeColor(black)
    c.line(mid + 118, box_top - 16, W - margin - 16, box_top - 16)
    c.setFont(TNRB, 9)
    c.drawString(mid + 8, box_top - 36, "Signature of Participant:")
    c.line(mid + 130, box_top - 38, W - margin - 16, box_top - 38)

    c.setFont(TNRB, 13)
    c.drawCentredString(W / 2, box_top - box_h - 20, title)

    table_top = box_top - box_h - 28
    table_bottom = margin + 28
    table_left = margin + 8
    table_right = W - margin - 8
    table_w = table_right - table_left
    row_h = (table_top - table_bottom) / 12.0

    c.setLineWidth(0.7)
    c.rect(table_left, table_bottom, table_w, table_top - table_bottom)
    for i in range(1, 12):
        y = table_top - i * row_h
        c.line(table_left, y, table_right, y)

    for i, (q, opts) in enumerate(questions):
        y_top = table_top - i * row_h
        y_bot = y_top - row_h
        qtext = f"{i + 1}. {q}"
        q_lines = wrap(c, qtext, TNRB, 8.2, table_w - 12)
        ty = y_top - 11
        c.setFont(TNRB, 8.2)
        for line in q_lines[:2]:
            c.drawString(table_left + 6, ty, line)
            ty -= 10
        c.setFont(TNR, 8)
        col_w = table_w / 4.0
        for j, opt in enumerate(opts):
            ox = table_left + 8 + j * col_w
            oy = y_bot + 10
            c.rect(ox, oy, 7, 7, stroke=1, fill=0)
            olines = wrap(c, opt, TNR, 7.4, col_w - 20)
            c.drawString(ox + 11, oy, olines[0])
            if len(olines) > 1:
                c.drawString(ox + 11, oy - 9, olines[1])

    c.setFont(TNR, 11)
    c.drawCentredString(W / 2, margin - 2, str(page_no))


def build_pdfs(logo):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(SURVEY_PDF), pagesize=letter)
    draw_page(c, title="Survey Questionnaire", questions=SURVEY, page_no=89, logo_path=logo)
    c.save()
    c = canvas.Canvas(str(FEEDBACK_PDF), pagesize=letter)
    draw_page(c, title="Feedback Questionnaire", questions=FEEDBACK, page_no=90, logo_path=logo)
    c.save()
    print("wrote", SURVEY_PDF, FEEDBACK_PDF)


def set_run(run, *, size=11, bold=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build_docx(path, title, questions, start):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.6)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    section._sectPr.append(pg)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nirmala Memorial Foundation College of Commerce and Science (Autonomous)")
    set_run(r, size=13, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Information Technology")
    set_run(r, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Community Engagement Project Questionnaire")
    set_run(r, size=11, bold=True)

    info = doc.add_table(rows=2, cols=2)
    info.style = "Table Grid"
    info.cell(0, 0).text = f"Name of Student(s): {STUDENTS}"
    info.cell(0, 1).text = "Name of Participant: ____________________"
    info.cell(1, 0).text = f"Seat No: {SEATS}"
    info.cell(1, 1).text = "Signature of Participant: ____________________"
    for row in info.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run(r, size=14, bold=True)

    tbl = doc.add_table(rows=12, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (q, opts) in enumerate(questions):
        cell = tbl.cell(i, 0)
        cell.text = ""
        pq = cell.paragraphs[0]
        rq = pq.add_run(f"{i + 1}. {q}")
        set_run(rq, size=10, bold=True)
        po = cell.add_paragraph()
        ro = po.add_run("     ".join("☐  " + o for o in opts))
        set_run(ro, size=9)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print("wrote", path)


APPS_SCRIPT = r'''/**
 * Open https://script.google.com with YOUR Google account.
 * New project → paste this file → Run createBothForms.
 * Grant permissions, then View → Logs for the live form links.
 */
function createBothForms() {
  var survey = buildForm_(
    "CEP Survey Questionnaire – IoT Bus Stop Overcrowding Prediction",
    SURVEY_ITEMS_
  );
  var feedback = buildForm_(
    "CEP Feedback Questionnaire – IoT Bus Stop Overcrowding Prediction",
    FEEDBACK_ITEMS_
  );
  Logger.log("SURVEY VIEW: " + survey.getPublishedUrl());
  Logger.log("SURVEY EDIT: " + survey.getEditUrl());
  Logger.log("FEEDBACK VIEW: " + feedback.getPublishedUrl());
  Logger.log("FEEDBACK EDIT: " + feedback.getEditUrl());
}

var DESC_ =
  "Nirmala Memorial Foundation College of Commerce and Science (Autonomous)\\n" +
  "Department of Information Technology\\n" +
  "Students: Mr. Om Keluskar (24TIT078), Mr. Jashith Agre (24TIT011)\\n" +
  "Project: IoT-Based Bus Stop System for Public Transport Overcrowding Prediction Using Machine Learning\\n" +
  "Site: Dindoshi Bus Depot, BEST Route 326";

function buildForm_(title, items) {
  var form = FormApp.create(title);
  form.setDescription(DESC_);
  form.setCollectEmail(false);
  form.setAllowResponseEdits(false);
  form.addTextItem().setTitle("Name of Participant").setRequired(true);
  items.forEach(function (it) {
    form.addMultipleChoiceItem()
      .setTitle(it.q)
      .setChoiceValues(it.opts)
      .setRequired(true);
  });
  return form;
}

var SURVEY_ITEMS_ = [
  {q:"How frequently do you experience overcrowding at Dindoshi Bus Depot or the BEST Route 326 bus stop?", opts:["Every Day","Several Times a Week","Occasionally","Rarely"]},
  {q:"Which mode of transport do you use most frequently to reach this stop?", opts:["BEST Bus (Route 326)","Other BEST Bus","Train then BEST Bus","Walking"]},
  {q:"What is the biggest problem you face because of overcrowding at the bus stop?", opts:["Missed bus / cannot board","Pushing or injury risk","Long uncomfortable wait","No crowding warning"]},
  {q:"Have you ever been unable to board a bus because the stop or doorway was too crowded?", opts:["Yes, Multiple Times","Yes, Once","No","Not Applicable"]},
  {q:"How safe do you feel while waiting at an overcrowded bus stop?", opts:["Very Safe","Safe","Unsafe","Very Unsafe"]},
  {q:"Have you ever reported overcrowding to depot staff or BEST officials?", opts:["Yes","No","Didn't Know How","Never Felt It Would Help"]},
  {q:"If you did report overcrowding, how satisfied were you with the response?", opts:["Very Satisfied","Satisfied","Dissatisfied","Never Reported"]},
  {q:"How important is it to warn commuters as soon as overcrowding is detected?", opts:["Extremely Important","Important","Slightly Important","Not Important"]},
  {q:"Do you think rain, heat, or peak hours make overcrowding at the stop more dangerous?", opts:["Strongly Agree","Agree","Disagree","Strongly Disagree"]},
  {q:"Would you trust an automated IoT system that predicts overcrowding using on-site lights, without requiring a mobile app?", opts:["Definitely Yes","Probably Yes","Probably No","Definitely No"]},
  {q:"Which feature would be most useful in a smart bus-stop overcrowding system?", opts:["On-site LED display","Telegram staff alerts","Live risk percentage","Bus-left operator button"]},
  {q:"Do you believe an IoT-based overcrowding prediction system can improve commuter safety at BEST bus stops?", opts:["Strongly Agree","Agree","Disagree","Strongly Disagree"]}
];

var FEEDBACK_ITEMS_ = [
  {q:"How easy was it to understand the purpose of the IoT bus-stop overcrowding system?", opts:["Very Easy","Easy","Neutral","Difficult"]},
  {q:"How useful do you think automatic passenger counting using the IoT sensor would be for overcrowding awareness?", opts:["Very Useful","Useful","Moderately Useful","Not Useful at All"]},
  {q:"How clear and understandable was the demonstrated overcrowding-prediction process?", opts:["Very Clear","Clear","Neutral","Unclear"]},
  {q:"How useful were the on-site red and green LED indicators for understanding crowding?", opts:["Very Useful","Useful","Moderately Useful","Not Useful at All"]},
  {q:"How useful would real-time Telegram alerts be for depot staff when crowding becomes CRITICAL?", opts:["Very Useful","Useful","Moderately Useful","Not Useful at All"]},
  {q:"How easy was it to understand the crowding information shown on the dashboard (count, weather, risk, status)?", opts:["Very Easy","Easy","Neutral","Difficult"]},
  {q:"How clearly did the hardware demonstration show how the sensor detects passenger crossings?", opts:["Very Clear","Clear","Neutral","Unclear"]},
  {q:"How confident are you that this system can warn of overcrowding faster than manual visual guessing?", opts:["Very Confident","Confident","Moderately Confident","Slightly Confident"]},
  {q:"Would you be willing to use this system if it remained available at Dindoshi Route 326?", opts:["Definitely Yes","Probably Yes","Not Sure","Definitely No"]},
  {q:"How clear and easy to understand were the NORMAL, WARNING, and CRITICAL status lights?", opts:["Very Clear","Clear","Neutral","Unclear"]},
  {q:"How useful do you think this system would be for improving crowding awareness at your bus stop?", opts:["Very Useful","Useful","Moderately Useful","Not Useful at All"]},
  {q:"Overall, how satisfied are you with the IoT bus-stop overcrowding system and its potential for commuter safety?", opts:["Very Satisfied","Satisfied","Neutral","Dissatisfied"]}
];
'''


def main():
    logo = prepare_logo()
    build_pdfs(logo)
    build_docx(SURVEY_DOCX, "Survey Questionnaire", SURVEY, 89)
    build_docx(FEEDBACK_DOCX, "Feedback Questionnaire", FEEDBACK, 90)
    script = ROOT / "scripts/Create_CEP_Google_Forms.gs"
    script.write_text(APPS_SCRIPT)
    print("wrote", script)


if __name__ == "__main__":
    main()
