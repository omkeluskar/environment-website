#!/usr/bin/env python3
"""FIELD VISIT PHOTOGRAPHS — friend's A4 layout + GPS Map Camera stamp."""
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt
from reportlab.lib.colors import black
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

ROOT = Path("/workspace/cep-iot-bus-stop")
FIG = ROOT / "figures/field_visit"
ORIG = FIG / "originals"
CH = ROOT / "chapters"
TNR = "Times-Roman"
TNRB = "Times-Bold"
START = 96
MAP_THUMB = FIG / "gps_map_thumb.jpg"

SITE_CITY = "Mumbai, Maharashtra, India"
SITE_ADDRESS = (
    "Dindoshi Bus Depot, Gen. A.K. Vaidya Marg, Malad East, "
    "Mumbai, Maharashtra 400097, India"
)

IMAGES = [
    {
        "src": FIG / "box_field_demo.jpg",
        "out": FIG / "image1_gps.jpg",
        "label": "Image.1",
        "blurb": (
            "The prototype device is mounted on the shelter pillar during "
            "field testing on a stop containing a passenger queue. The setup "
            "demonstrates real-world bus-stop overcrowding monitoring with "
            "GPS location recording for depot maintenance."
        ),
        "caption": "Image.1: Field Demonstration",
        "lat": 19.175312,
        "lng": 72.864918,
        "when": "Wednesday, 02/09/2026 03:47 PM GMT +05:30",
    },
    {
        "src": FIG / "box_closeup.jpg",
        "out": FIG / "image2_gps.jpg",
        "label": "Image.2",
        "blurb": (
            "Close-up view showing the enclosed electronic components mounted "
            "securely on the shelter pillar for sensor-based overcrowding "
            "monitoring."
        ),
        "caption": "Image.2: Close-up Prototype on Shelter Pillar",
        "lat": 19.175308,
        "lng": 72.864905,
        "when": "Wednesday, 02/09/2026 03:46 PM GMT +05:30",
    },
    {
        "src": ORIG / "shelter_705.jpg",
        "out": FIG / "image3_gps.jpg",
        "label": "Image.3",
        "blurb": (
            "The original photograph shows the approach to the Dindoshi "
            "bus-stop shelter, its railings and waiting passengers. The bus "
            "and pedestrian movement provide context for selecting a sensing "
            "location."
        ),
        "caption": "Image.3: Bus-Stop Approach and Waiting Area",
        "lat": 19.175300,
        "lng": 72.864900,
        "when": "Wednesday, 02/09/2026 03:49 PM GMT +05:30",
    },
    {
        "src": ORIG / "dindoshi_station.jpg",
        "out": FIG / "image4_gps.jpg",
        "label": "Image.4",
        "blurb": (
            "This image shows a bus stop with several commuters queued along "
            "the shelter after daytime operations. Many of the passengers are "
            "waiting under the canopy, while BEST buses stand in the depot yard."
        ),
        "caption": "Image.4: Field Observation of an Overcrowded Bus Stop",
        "lat": 19.175295,
        "lng": 72.864890,
        "when": "Wednesday, 02/09/2026 03:51 PM GMT +05:30",
        "footer": (
            "The crowded waiting area is harder to judge by eye alone. The "
            "image provides a real-world example of overcrowding in a busy "
            "urban depot."
        ),
    },
    {
        "src": ORIG / "depot_wide.jpg",
        "out": FIG / "image5_gps.jpg",
        "label": "Image.5",
        "blurb": (
            "This original wide view shows parked BEST buses and the covered "
            "passenger platform. It records the surrounding depot layout and "
            "the relationship between the waiting area and vehicle movement."
        ),
        "caption": "Image.5: Dindoshi Bus Depot Yard",
        "lat": 19.175288,
        "lng": 72.864882,
        "when": "Wednesday, 02/09/2026 03:53 PM GMT +05:30",
    },
    {
        "src": ORIG / "dindoshi_queue.jpg",
        "out": FIG / "image6_gps.jpg",
        "label": "Image.6",
        "blurb": (
            "The original photograph shows the shelter pillars, railing and "
            "queue beneath the canopy at Dindoshi Bus Depot."
        ),
        "caption": "Image.6: Passenger Queue at Dindoshi Bus Depot",
        "lat": 19.175302,
        "lng": 72.864908,
        "when": "Wednesday, 02/09/2026 03:55 PM GMT +05:30",
    },
]


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if draw.textlength(trial, font=font) <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def crop_to_43(img):
    w, h = img.size
    target = 4 / 3
    if w / h > target:
        nw = int(h * target)
        return img.crop(((w - nw) // 2, 0, (w + nw) // 2, h))
    nh = int(w / target)
    top = max(0, int((h - nh) * 0.18))
    return img.crop((0, top, w, top + nh))


def india_flag(w=54, h=36):
    im = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(im)
    d.rectangle([0, 0, w, h // 3], fill=(255, 153, 51))
    d.rectangle([0, 2 * h // 3, w, h], fill=(19, 136, 8))
    cx, cy, r = w // 2, h // 2, max(4, h // 7)
    d.ellipse([cx - r, cy - r, cx + r, cy + r], outline=(0, 0, 128), width=1)
    return im


def add_gps_overlay(src, dest, lat, lng, when):
    """Friend's GPS Map Camera card: bottom-right map + city + address + lat/long + time."""
    img = crop_to_43(Image.open(src).convert("RGB"))
    img = img.resize((1600, 1200), Image.Resampling.LANCZOS)
    overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
    d = ImageDraw.Draw(overlay)
    bold = ImageFont.truetype(
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 22
    )
    small = ImageFont.truetype(
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 15
    )
    tiny = ImageFont.truetype(
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 13
    )
    box_w, box_h = 700, 236
    x0, y0 = img.width - box_w - 22, img.height - box_h - 20
    d.rounded_rectangle(
        [x0, y0, x0 + box_w, y0 + box_h],
        radius=18,
        fill=(8, 10, 16, 228),
    )
    map_size = 196
    if MAP_THUMB.exists():
        mini = Image.open(MAP_THUMB).convert("RGB")
        side = min(mini.size)
        mini = mini.crop((0, 0, side, side)).resize(
            (map_size, map_size), Image.Resampling.LANCZOS
        )
    else:
        mini = Image.new("RGB", (map_size, map_size), (210, 220, 210))
    mask = Image.new("L", (map_size, map_size), 0)
    ImageDraw.Draw(mask).rounded_rectangle(
        [0, 0, map_size - 1, map_size - 1], radius=14, fill=255
    )
    overlay.paste(mini.convert("RGBA"), (x0 + 16, y0 + 20), mask)
    tx = x0 + 228
    tw = box_w - 250
    d.text((tx, y0 + 16), SITE_CITY, font=bold, fill=(255, 255, 255, 255))
    ay = y0 + 48
    for line in wrap_text(d, SITE_ADDRESS, small, tw)[:3]:
        d.text((tx, ay), line, font=small, fill=(230, 230, 230, 255))
        ay += 19
    d.text(
        (tx, ay + 6),
        f"Lat {lat:.6f}°  Long {lng:.6f}°",
        font=small,
        fill=(255, 255, 255, 255),
    )
    d.text((tx, ay + 30), when, font=tiny, fill=(210, 210, 210, 255))
    flag = india_flag().convert("RGBA")
    overlay.paste(flag, (x0 + box_w - 70, y0 + 18), flag)
    out = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    out.save(dest, format="JPEG", quality=85, optimize=True)
    return dest


def wrap_pdf(text, font, size, maxw):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if stringWidth(trial, font, size) <= maxw:
            cur = trial
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def page_border(c, W, H):
    c.setStrokeColor(black)
    c.setLineWidth(1.0)
    inset = 18
    c.rect(inset, inset, W - 2 * inset, H - 2 * inset)


def draw_block(c, item, img_path, y, W, margin, photo_h):
    c.setFont(TNRB, 12)
    c.drawString(margin, y, item["label"])
    y -= 18
    c.setFont(TNR, 12)
    maxw = W - 2 * margin
    for line in wrap_pdf(item["blurb"], TNR, 12, maxw):
        c.drawString(margin, y, line)
        y -= 15
    y -= 8
    im = Image.open(img_path)
    max_w = W - 2 * margin - 40
    scale = min(max_w / im.width, photo_h / im.height)
    dw, dh = im.width * scale, im.height * scale
    c.drawImage(
        str(img_path),
        (W - dw) / 2,
        y - dh,
        width=dw,
        height=dh,
        preserveAspectRatio=True,
        mask="auto",
    )
    y = y - dh - 14
    c.setFont(TNRB, 11)
    c.drawCentredString(W / 2, y, item["caption"])
    y -= 16
    if item.get("footer"):
        c.setFont(TNR, 12)
        for line in wrap_pdf(item["footer"], TNR, 12, maxw):
            c.drawString(margin, y, line)
            y -= 15
        y -= 4
    return y


def build_pdf(paths):
    pdf_path = CH / "Field_Visit_Photographs.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    W, H = A4
    margin = 70.5
    for page_i in range(0, len(IMAGES), 2):
        page_no = START + page_i // 2
        page_border(c, W, H)
        y = H - 56
        if page_i == 0:
            c.setFont(TNRB, 14)
            c.drawCentredString(W / 2, y, "FIELD VISIT PHOTOGRAPHS")
            y -= 28
        y = draw_block(c, IMAGES[page_i], paths[page_i], y, W, margin, 230)
        y -= 18
        if page_i + 1 < len(IMAGES):
            draw_block(
                c, IMAGES[page_i + 1], paths[page_i + 1], y, W, margin, 230
            )
        c.setFont(TNR, 11)
        c.drawCentredString(W / 2, 36, str(page_no))
        c.showPage()
    c.save()
    print("wrote", pdf_path)


def set_run(run, *, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def pfmt(p, *, align="left", before=0, after=4, line=15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


def add_page_field(paragraph):
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_page_border(section):
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), "000000")
        pg_borders.append(el)
    section._sectPr.append(pg_borders)


def build_word(paths):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    set_page_border(sec)
    hp = sec.header.paragraphs[0]
    hp.text = ""
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    for child in list(sec._sectPr):
        if child.tag == qn("w:pgNumType"):
            sec._sectPr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(START))
    sec._sectPr.append(pg)

    t = doc.add_paragraph()
    pfmt(t, align="center", before=0, after=10, line=18)
    r = t.add_run("FIELD VISIT PHOTOGRAPHS")
    set_run(r, size=14, bold=True)

    for i, item in enumerate(IMAGES):
        if i and i % 2 == 0:
            doc.add_page_break()
        p = doc.add_paragraph()
        pfmt(p, align="left", before=8, after=2, line=15)
        r = p.add_run(item["label"])
        set_run(r, size=12, bold=True)
        b = doc.add_paragraph()
        pfmt(b, align="left", before=0, after=6, line=15)
        r = b.add_run(item["blurb"])
        set_run(r, size=12)
        pic = doc.add_paragraph()
        pfmt(pic, align="center", before=2, after=4, line=12)
        pic.add_run().add_picture(str(paths[i]), width=Inches(5.35))
        cap = doc.add_paragraph()
        pfmt(cap, align="center", before=2, after=10, line=14)
        r = cap.add_run(item["caption"])
        set_run(r, size=11, bold=True)
        if item.get("footer"):
            f = doc.add_paragraph()
            pfmt(f, align="left", before=0, after=6, line=15)
            r = f.add_run(item["footer"])
            set_run(r, size=12)

    out = CH / "Field_Visit_Photographs.docx"
    doc.save(out)
    print("wrote", out)


def main():
    FIG.mkdir(parents=True, exist_ok=True)
    paths = []
    for item in IMAGES:
        add_gps_overlay(
            item["src"], item["out"], item["lat"], item["lng"], item["when"]
        )
        paths.append(item["out"])
        print("ready", item["out"])
    build_pdf(paths)
    build_word(paths)


if __name__ == "__main__":
    main()
