#!/usr/bin/env python3
"""Keep the original college form pages. Only strip overlay images and
align questionnaire checkboxes with their option labels."""
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path("/workspace/cep-iot-bus-stop")
SRC = ROOT / "figures/front_matter/source_forms.pdf"
CH = ROOT / "chapters"
TMP = Path("/tmp/cep_form_pages")


def strip_large_overlays(page, min_w=400, min_h=80):
    for info in page.get_images(full=True):
        xref, w, h = info[0], info[2], info[3]
        if w >= min_w and h >= min_h:
            page.delete_image(xref)


def _spans(page):
    out = []
    for b in page.get_text("dict")["blocks"]:
        if b.get("type") != 0:
            continue
        for line in b["lines"]:
            for s in line["spans"]:
                t = s["text"]
                if not t.strip():
                    continue
                out.append(s)
    out.sort(key=lambda s: (round(s["bbox"][1], 1), s["bbox"][0]))
    return out


def align_checkboxes(page):
    """Place a square immediately left of each option label; drop stray ☐."""
    spans = _spans(page)
    option_boxes = []
    checkbox_boxes = []
    for s in spans:
        text = s["text"].strip()
        font = s.get("font", "")
        if "FreeMono" in font or text == "☐":
            checkbox_boxes.append(pymupdf.Rect(s["bbox"]))
            continue
        if s["size"] >= 12:
            continue
        if s.get("flags", 0) & 2**4:  # superscript-ish; ignore
            pass
        # option labels are 10pt regular Times, not numbered questions
        bold = "Bold" in font
        if bold:
            continue
        if text[0].isdigit() and (len(text) == 1 or text[1] in ".)"):
            continue
        y0 = s["bbox"][1]
        if y0 < 240 or y0 > 720:
            continue
        if text.isdigit():
            continue
        option_boxes.append(s)

    for r in checkbox_boxes:
        shrink = pymupdf.Rect(r.x0 + 0.8, r.y0 + 0.8, r.x1 - 0.8, r.y1 - 0.8)
        page.add_redact_annot(shrink, fill=(1, 1, 1))
    if checkbox_boxes:
        page.apply_redactions(images=0)

    for s in option_boxes:
        x0, y0, x1, y1 = s["bbox"]
        h = y1 - y0
        box = 8.0
        gap = 4.5
        bx = x0 - box - gap
        by = y0 + (h - box) / 2.0
        if bx < 40:
            continue
        page.draw_rect(
            pymupdf.Rect(bx, by, bx + box, by + box),
            color=(0, 0, 0),
            fill=None,
            width=0.8,
        )


def write_page(src_doc, index, dest, *, checkboxes=False, strip=True):
    one = pymupdf.open()
    one.insert_pdf(src_doc, from_page=index, to_page=index)
    page = one[0]
    if strip:
        strip_large_overlays(page)
    if checkboxes:
        align_checkboxes(page)
    dest.parent.mkdir(parents=True, exist_ok=True)
    one.save(dest, deflate=True, garbage=4)
    one.close()
    print("wrote", dest)


def pdf_to_word(pdf_path, docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(0)
    sec.right_margin = Inches(0)
    sec.top_margin = Inches(0)
    sec.bottom_margin = Inches(0)
    pdf = pymupdf.open(pdf_path)
    TMP.mkdir(parents=True, exist_ok=True)
    for i, page in enumerate(pdf):
        if i:
            doc.add_page_break()
        pix = page.get_pixmap(matrix=pymupdf.Matrix(300 / 72, 300 / 72), alpha=False)
        png = TMP / f"{pdf_path.stem}_p{i + 1}.png"
        pix.save(png)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        run = p.add_run()
        run.add_picture(str(png), width=Inches(8.5), height=Inches(11))
    pdf.close()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print("wrote", docx_path)


def merge_front_matter():
    out = CH / "Front_Matter.pdf"
    merged = pymupdf.open()
    for p in (
        CH / "Cover_Page.pdf",
        CH / "Certificate_of_Project_Completion.pdf",
        CH / "Project_Proposal_Approval_Sheet.pdf",
        CH / "Table_of_Contents.pdf",
        CH / "CEP_Survey_Questionnaire.pdf",
        CH / "CEP_Feedback_Questionnaire.pdf",
    ):
        if p.exists():
            merged.insert_pdf(pymupdf.open(p))
    merged.save(out, deflate=True, garbage=4)
    print("wrote", out)


def main():
    src = pymupdf.open(SRC)
    # page 3 = proposal, 4 = survey, 5 = feedback (0-based 2,3,4)
    write_page(src, 2, CH / "Project_Proposal_Approval_Sheet.pdf", strip=True)
    write_page(src, 3, CH / "CEP_Survey_Questionnaire.pdf", strip=False)
    write_page(src, 4, CH / "CEP_Feedback_Questionnaire.pdf", strip=True, checkboxes=True)
    src.close()
    pdf_to_word(CH / "Project_Proposal_Approval_Sheet.pdf", CH / "Project_Proposal_Approval_Sheet.docx")
    pdf_to_word(CH / "CEP_Survey_Questionnaire.pdf", CH / "CEP_Survey_Questionnaire.docx")
    pdf_to_word(CH / "CEP_Feedback_Questionnaire.pdf", CH / "CEP_Feedback_Questionnaire.docx")
    merge_front_matter()


if __name__ == "__main__":
    main()
